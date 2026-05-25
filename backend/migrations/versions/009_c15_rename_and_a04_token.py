"""Rename C15 label to 'Important Notes & Special Conditions'; patch master DOCX.

Two fixes:
1. DB: renames master_fields label/appendix_row_label for C15 from
   'Optional Terms' to 'Important Notes & Special Conditions'.
2. Master DOCX: replaces the hard-coded 'Marina Tower Phase II' placeholder
   in the Appendix 'The Project' cell with the {{A04}} token, and renames
   the C15 row label from 'Optional Terms' to
   'Important Notes & Special Conditions'.

Revision ID: 009_c15_rename_and_a04_token
Revises: 008_new_roles_and_project_users
Create Date: 2026-05-25
"""

from pathlib import Path

import sqlalchemy as sa
from alembic import op

revision: str = "009_c15_rename_and_a04_token"
down_revision: str | None = "008_new_roles_and_project_users"
branch_labels = None
depends_on = None

_NEW_LABEL = "Important Notes & Special Conditions"
_OLD_LABEL = "Optional Terms"


def upgrade() -> None:
    # 1. Update DB labels
    op.execute(
        sa.text(
            """
            UPDATE master_fields
            SET    field_label       = :new_label,
                   appendix_row_label = :new_label
            WHERE  field_id = 'C15'
              AND  (field_label = :old_label OR appendix_row_label = :old_label)
            """
        ).bindparams(new_label=_NEW_LABEL, old_label=_OLD_LABEL)
    )

    # 2. Patch the master DOCX in-place (idempotent)
    _patch_master_docx()


def downgrade() -> None:
    op.execute(
        sa.text(
            """
            UPDATE master_fields
            SET    field_label       = :old_label,
                   appendix_row_label = :old_label
            WHERE  field_id = 'C15'
            """
        ).bindparams(old_label=_OLD_LABEL)
    )


# ---------------------------------------------------------------------------
# Master DOCX patch helpers
# ---------------------------------------------------------------------------

def _patch_master_docx() -> None:
    """Apply two fixes to the master DOCX:

    * Rename the C15 row label cell from 'Optional Terms' to the new label.
    * Replace the hard-coded 'Marina Tower Phase II' value in the Appendix
      'The Project' row with the {{A04}} token so the project name renders
      dynamically at PDF generation time.
    """
    try:
        from copy import deepcopy

        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        master_path = Path(__file__).resolve().parents[2] / "masters" / "sca_master_v1.docx"
        if not master_path.exists():
            print(f"  [009] master docx not found at {master_path} — skipping DOCX patch")
            return

        doc = Document(str(master_path))
        changed = False

        for table in doc.tables:
            for row in table.rows:
                cells = row.cells

                # Fix 1: rename C15 label cell
                if cells[0].text.strip() == _OLD_LABEL:
                    _set_cell_text(cells[0], _NEW_LABEL)
                    changed = True
                    print(f"  [009] DOCX: renamed C15 label → '{_NEW_LABEL}'")

                # Fix 2: replace hard-coded project name with {{A04}} token
                # The value cell is cells[2] for a 3-column table, else cells[-1]
                value_cell = cells[2] if len(cells) > 2 else cells[-1]
                label_txt = cells[0].text.strip()
                val_txt = value_cell.text.strip()
                if label_txt == "The Project" and val_txt not in ("{{A04}}", ""):
                    _set_cell_text(value_cell, "{{A04}}")
                    changed = True
                    print(f"  [009] DOCX: replaced '{val_txt[:40]}' with {{{{A04}}}} in 'The Project' cell")

        if changed:
            doc.save(str(master_path))
            print("  [009] master DOCX saved.")
        else:
            print("  [009] master DOCX already up-to-date — no change.")

    except Exception as exc:
        # Never block migrations on a DOCX patch failure
        print(f"  [009] WARNING: DOCX patch failed ({exc}); master DOCX unchanged.")


def _set_cell_text(cell, text: str) -> None:
    """Replace cell content with `text`, preserving the first run's formatting."""
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)
