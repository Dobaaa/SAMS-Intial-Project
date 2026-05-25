"""Replace A21 percentage input with A21_DISPLAY synthetic token for Max LD appendix row.

The 'Maximum Liquidated Damages' appendix row previously showed
'{{A21}}% of the Contract Price / {{A21_AMOUNT}} AED'. It now renders as
'AED X per day and limited to 10% of Subcontract Value' via the A21_DISPLAY
synthetic token (computed from C11 daily rate at PDF render time).

Changes:
1. DB: set A21 show_in_appendix=False (no longer a user-visible appendix field).
2. Master DOCX: change 'Maximum Liquidated Damages' value cell to {{A21_DISPLAY}}.
   (Belt-and-suspenders — the committed DOCX already carries this patch.)

Revision ID: 011_max_ld_display
Revises: 010_advance_payment_pct
Create Date: 2026-05-25
"""

from pathlib import Path

import sqlalchemy as sa
from alembic import op

revision: str = "011_max_ld_display"
down_revision: str | None = "010_advance_payment_pct"
branch_labels = None
depends_on = None


def upgrade() -> None:
    # Hide A21 from the appendix review tab — it's now a derived/synthetic value.
    op.execute(sa.text(
        "UPDATE master_fields SET show_in_appendix = false WHERE field_id = 'A21'"
    ))
    _patch_master_docx()


def downgrade() -> None:
    op.execute(sa.text(
        "UPDATE master_fields SET show_in_appendix = true WHERE field_id = 'A21'"
    ))


def _patch_master_docx() -> None:
    try:
        from docx import Document

        master_path = Path(__file__).resolve().parents[2] / "masters" / "sca_master_v1.docx"
        if not master_path.exists():
            print(f"  [011] master docx not found — skipping")
            return

        doc = Document(str(master_path))
        changed = False
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if cells[0].text.strip() == "Maximum Liquidated Damages":
                    val_cell = cells[2] if len(cells) > 2 else cells[-1]
                    if "A21" in val_cell.text and "A21_DISPLAY" not in val_cell.text:
                        _set_cell_text(val_cell, "{{A21_DISPLAY}}")
                        changed = True
                        print("  [011] DOCX: Maximum Liquidated Damages → {{A21_DISPLAY}}")
        if changed:
            doc.save(str(master_path))
            print("  [011] master DOCX saved.")
        else:
            print("  [011] master DOCX already up-to-date.")
    except Exception as exc:
        print(f"  [011] WARNING: DOCX patch failed ({exc})")


def _set_cell_text(cell, text: str) -> None:
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
