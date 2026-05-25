"""Add C03_PCT (Advance Payment %) field + update Appendix row to show 'AED X as Y%'.

Changes:
1. DB: insert C03_PCT into all conditions master templates, default="10".
2. Master DOCX: change Appendix 'Advance Payment Amount' value cell from
   {{C03}} to {{C03_DISPLAY}} (a synthetic token rendered at PDF time as
   "AED {amount} as {pct}% of the Subcontract Price").

Revision ID: 010_advance_payment_pct
Revises: 009_c15_rename_and_a04_token
Create Date: 2026-05-25
"""

from pathlib import Path

import sqlalchemy as sa
from alembic import op

revision: str = "010_advance_payment_pct"
down_revision: str | None = "009_c15_rename_and_a04_token"
branch_labels = None
depends_on = None


def upgrade() -> None:
    # 1. Add C03_PCT field to all conditions templates (idempotent via NOT EXISTS)
    op.execute(sa.text("""
        INSERT INTO master_fields (
            id, template_id, field_id, clause_number,
            field_label, input_type, is_required,
            default_value, show_in_appendix, sort_order
        )
        SELECT
            gen_random_uuid(), mt.id, 'C03_PCT', '3.4.1',
            'Advance Payment (%)', 'number', false,
            '10', false, 3
        FROM master_templates mt
        WHERE mt.type = 'conditions'
          AND NOT EXISTS (
            SELECT 1 FROM master_fields mf2
            WHERE mf2.template_id = mt.id AND mf2.field_id = 'C03_PCT'
          )
    """))

    # 2. Patch master DOCX: {{C03}} → {{C03_DISPLAY}} in Appendix row
    _patch_master_docx()


def downgrade() -> None:
    op.execute(sa.text("DELETE FROM master_fields WHERE field_id = 'C03_PCT'"))


def _patch_master_docx() -> None:
    try:
        from docx import Document

        master_path = Path(__file__).resolve().parents[2] / "masters" / "sca_master_v1.docx"
        if not master_path.exists():
            print(f"  [010] master docx not found — skipping DOCX patch")
            return

        doc = Document(str(master_path))
        changed = False

        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if cells[0].text.strip() == "Advance Payment Amount":
                    val_cell = cells[2] if len(cells) > 2 else cells[-1]
                    if val_cell.text.strip() == "{{C03}}":
                        _set_cell_text(val_cell, "{{C03_DISPLAY}}")
                        changed = True
                        print("  [010] DOCX: Advance Payment cell → {{C03_DISPLAY}}")

        if changed:
            doc.save(str(master_path))
            print("  [010] master DOCX saved.")
        else:
            print("  [010] master DOCX already up-to-date.")

    except Exception as exc:
        print(f"  [010] WARNING: DOCX patch failed ({exc})")


def _set_cell_text(cell, text: str) -> None:
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
