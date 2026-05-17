"""Clause extraction + per-agreement revision CRUD.

The "clauses" extracted here are the paragraphs of the master docx that
admin can revise on a per-agreement basis. Each clause is anchored by a
SHA-256 hash of its trimmed/normalised text so revisions survive minor
re-tokenization of the master.

v2.0 — silent application of accepted revisions during render.
v2.1 will add the accept/reject workflow + role gating.
v2.2 will surface pending revisions as Word track-changes in the PDF.
"""
from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.text.paragraph import Paragraph

from services.docx_pdf_service import MASTER_DOCX, TOKEN_RE

# Paragraphs shorter than this are too small to be meaningful clauses (page
# numbers, single tokens, empty lines, etc.) — skip them in the picker.
MIN_CLAUSE_CHAR_LENGTH = 12

# Cached clause list keyed by master docx mtime. Rebuilt automatically on
# every re-tokenization since mtime changes.
_CLAUSE_CACHE: dict[float, list["MasterClause"]] = {}


@dataclass(frozen=True, slots=True)
class MasterClause:
    """One revisable paragraph in the master docx."""

    clause_hash: str
    clause_label: str
    text: str          # As stored in the docx, with {{FIELD_ID}} tokens intact
    section: str       # Best-effort section ("Form", "Conditions", "Appendix", "Cover")
    position: int      # Order in the document (informational; not used for matching)

    def as_dict(self) -> dict:
        return {
            "clause_hash": self.clause_hash,
            "clause_label": self.clause_label,
            "text": self.text,
            "section": self.section,
            "position": self.position,
        }


def _normalise(text: str) -> str:
    """Whitespace-normalise a paragraph for hashing."""
    return re.sub(r"\s+", " ", text).strip()


def _hash_clause(text: str) -> str:
    """SHA-256 of the normalised paragraph text. Used as the stable anchor."""
    return hashlib.sha256(_normalise(text).encode("utf-8")).hexdigest()


def _label_for_paragraph(text: str, section: str, position: int) -> str:
    """Pick a short human label for the clause picker UI.

    First non-whitespace stretch up to ~80 chars; prefix with section so
    the list stays scannable. Numbered clause prefixes (e.g. "1.3
    Priority of Documents") naturally fall through and become the label.
    """
    stripped = _normalise(text)
    head = stripped[:80]
    if len(stripped) > 80:
        head = head.rsplit(" ", 1)[0] + "…"
    return f"{section} · {head}" if section else head


def _detect_section(text: str, current_section: str) -> str:
    """Update the running 'section' header based on paragraph content.

    The master docx has narrative paragraphs that name each major
    section (e.g. "Form of Subcontract Agreement", "Conditions of the
    Subcontract Agreement"). This helps us tag each clause without
    needing the docx's style metadata.
    """
    norm = _normalise(text).lower()
    if "form of subcontract agreement" == norm or norm == "subcontract agreement":
        return "Form"
    if "appendix to the subcontract" in norm:
        return "Appendix"
    if "conditions of the subcontract" in norm:
        return "Conditions"
    if "annexures of the subcontract" in norm:
        return "Annexures"
    return current_section


def _iter_doc_paragraphs(doc) -> Iterable[tuple[Paragraph, str]]:
    """Yield (paragraph, section) tuples in document order — body paragraphs
    plus table-cell paragraphs."""
    section = "Cover"
    for p in doc.paragraphs:
        section = _detect_section(p.text, section)
        yield p, section
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    # Tables sit in either the Form (signature blocks),
                    # Appendix (Information & Data tables), or the
                    # Annexures section in this master. Pragmatically
                    # tag the entire table-paragraph batch as "Appendix"
                    # since the signature block paragraphs are usually
                    # filtered out by MIN_CLAUSE_CHAR_LENGTH.
                    yield p, "Appendix"


def list_master_clauses(master_path: Path | None = None) -> list[MasterClause]:
    """Walk the master docx and produce a stable list of revisable clauses.

    Cached by docx mtime; recomputes after re-tokenization.
    """
    path = master_path or MASTER_DOCX
    if not path.exists():
        return []
    mtime = path.stat().st_mtime
    cached = _CLAUSE_CACHE.get(mtime)
    if cached is not None:
        return cached

    doc = Document(str(path))
    clauses: list[MasterClause] = []
    seen_hashes: set[str] = set()
    for idx, (para, section) in enumerate(_iter_doc_paragraphs(doc)):
        text = para.text or ""
        norm = _normalise(text)
        if len(norm) < MIN_CLAUSE_CHAR_LENGTH:
            continue
        # Skip lines that are JUST a {{FIELD_ID}} token — they're handled
        # by the field editor on the Document view, not the clause editor.
        if TOKEN_RE.sub("", norm).strip() == "":
            continue
        clause_hash = _hash_clause(text)
        if clause_hash in seen_hashes:
            # Deduplicate paragraphs whose normalised text is identical
            # (e.g. repeated boilerplate). First occurrence wins; revisions
            # to one will be re-applied to all matches at render time.
            continue
        seen_hashes.add(clause_hash)
        clauses.append(
            MasterClause(
                clause_hash=clause_hash,
                clause_label=_label_for_paragraph(text, section, idx),
                text=text,
                section=section,
                position=idx,
            )
        )

    _CLAUSE_CACHE.clear()
    _CLAUSE_CACHE[mtime] = clauses
    return clauses


def find_master_clause_by_hash(clause_hash: str) -> MasterClause | None:
    for c in list_master_clauses():
        if c.clause_hash == clause_hash:
            return c
    return None


def apply_accepted_revisions_to_doc(doc, revisions: Iterable[tuple[str, str]]) -> int:
    """Walk every paragraph and replace its text with the modified_text
    whenever its normalised text hashes to one of the accepted revisions.

    `revisions` is an iterable of ``(clause_hash, modified_text)`` pairs.
    Returns the count of paragraphs that were rewritten.
    """
    rev_map = {h: m for h, m in revisions}
    if not rev_map:
        return 0
    replaced = 0
    for para in _iter_doc_paragraphs_for_apply(doc):
        h = _hash_clause(para.text or "")
        modified = rev_map.get(h)
        if modified is None:
            continue
        _rewrite_paragraph(para, modified)
        replaced += 1
    return replaced


def _iter_doc_paragraphs_for_apply(doc):
    """Same as _iter_doc_paragraphs but yields only the paragraph (no
    section), and traverses tables too."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _rewrite_paragraph(para: Paragraph, new_text: str) -> None:
    """Replace paragraph text with `new_text`, collapsing into the first run.

    Loses per-run inline formatting WITHIN the replaced span — acceptable
    because clauses are plain prose, and the paragraph-level formatting
    (alignment, indent, list numbering) lives on the paragraph properties
    and is unaffected by this run-content swap.
    """
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""


# OOXML track-changes namespace constants. <w:del>/<w:ins> are the
# standard Word revision elements; LibreOffice renders them with
# strikethrough (delete) and single-underline (insert) inline.
_WORDML_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def apply_pending_revisions_to_doc(
    doc,
    revisions: Iterable[tuple[str, str, str]],
    *,
    author: str = "SAMS Reviewer",
    when: datetime | None = None,
) -> int:
    """Wrap each matching paragraph in OOXML track-changes markup.

    ``revisions`` is an iterable of ``(clause_hash, original_text, modified_text)``
    tuples — typically one per pending ``AgreementClauseRevision``. For
    each match, the paragraph's existing content (which after upstream
    accepted-revision processing already equals ``original_text``) is
    rewrapped as a ``<w:del>`` block, and the modified text is appended
    as a ``<w:ins>`` block. LibreOffice then renders the strikethrough +
    underline track-changes pair when the PDF is generated.

    Returns the count of paragraphs that were marked.
    """
    from copy import deepcopy
    from lxml import etree

    rev_map = {h: (orig, mod) for h, orig, mod in revisions}
    if not rev_map:
        return 0

    when_iso = (when or _now()).strftime("%Y-%m-%dT%H:%M:%SZ")
    NS = "{" + _WORDML_NS + "}"
    marked = 0
    next_id = 1000  # avoid colliding with any pre-existing w:id values in the master

    for para in _iter_doc_paragraphs_for_apply(doc):
        h = _hash_clause(para.text or "")
        if h not in rev_map:
            continue
        _, modified_text = rev_map[h]
        original_text = para.text or ""

        p_el = para._element
        # Preserve the paragraph properties (alignment, list numbering,
        # indent) AND the formatting of the first run so the track-changes
        # spans inherit the same font/size/color as the surrounding clause.
        pPr = p_el.find(f"{NS}pPr")
        first_r = p_el.find(f"{NS}r")
        rPr_template = first_r.find(f"{NS}rPr") if first_r is not None else None

        # Wipe the paragraph's current children except pPr; we'll rebuild
        # below with the track-changes structure.
        for child in list(p_el):
            if child is pPr:
                continue
            p_el.remove(child)

        # <w:del id=".." author=".." date=".."> -> <w:r>[<w:rPr/>]<w:delText/></w:r></w:del>
        del_el = etree.SubElement(
            p_el,
            f"{NS}del",
            {
                f"{NS}id": str(next_id),
                f"{NS}author": author,
                f"{NS}date": when_iso,
            },
        )
        r_del = etree.SubElement(del_el, f"{NS}r")
        if rPr_template is not None:
            r_del.append(deepcopy(rPr_template))
        del_text = etree.SubElement(r_del, f"{NS}delText")
        del_text.text = original_text
        del_text.set(_XML_SPACE, "preserve")
        next_id += 1

        # <w:ins ...> -> <w:r>[<w:rPr/>]<w:t/></w:r></w:ins>
        ins_el = etree.SubElement(
            p_el,
            f"{NS}ins",
            {
                f"{NS}id": str(next_id),
                f"{NS}author": author,
                f"{NS}date": when_iso,
            },
        )
        r_ins = etree.SubElement(ins_el, f"{NS}r")
        if rPr_template is not None:
            r_ins.append(deepcopy(rPr_template))
        ins_text = etree.SubElement(r_ins, f"{NS}t")
        ins_text.text = modified_text
        ins_text.set(_XML_SPACE, "preserve")
        next_id += 1

        marked += 1

    return marked


def _now() -> datetime:
    return datetime.now(UTC)
