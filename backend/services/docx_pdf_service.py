"""DOCX-based PDF rendering for SCA agreements.

Uses the BGCC-tuned master docx (backend/masters/sca_master_v1.docx) as
the source template. The master is already arranged so LibreOffice
renders it to exactly 42 pages (matching the BGCC source PDF) on a
system with Tahoma + Times New Roman + Arial + Georgia + Caladea
installed. Each agreement-specific spot in the master carries a
``{{FIELD_ID}}`` token (``{{F02}}``, ``{{C03}}``, ``{{A07}}``, etc.) that
this service substitutes with the agreement's stored values at render
time, then hands the populated docx to ``libreoffice --headless --convert-to
pdf``.

Token rules
-----------
* ``{{FIELD_ID}}`` — substituted with the agreement's stored value for
  that field. HTML escaping is unnecessary inside a docx (Word handles
  the encoding internally).
* If the field's input type is ``date``, the value is reformatted via
  ``_format_longdate`` to match Rev 01 item 1 ("05th May 2026").
* Tokens whose field is empty or unknown are replaced with an empty
  string (the surrounding text stays — e.g. "AED" suffix in an appendix
  cell stays put).
* Tokens that DON'T appear in the master are silently ignored.

Adding a new field to the master
--------------------------------
Edit ``backend/masters/sca_master_v1.docx`` in Word/LibreOffice and type
``{{F99}}`` (or whichever ID) anywhere the value should appear. Save.
That's it — no code change needed here.
"""
from __future__ import annotations

import logging
import re
import subprocess
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTERS_DIR = BACKEND_DIR / "masters"
MASTER_DOCX = MASTERS_DIR / "sca_master_v1.docx"

# Field IDs whose values should be reformatted as a long-form date
# ("05th May 2026") at substitution time. The DB stores them as ISO
# strings; the master's tokens don't know that.
DATE_FIELDS: set[str] = {"F01"}

# Field IDs whose values are stored as raw numbers but should render
# with thousands separators + 2-dp decimals per Rev 01 items 11/12
# (e.g. F08 "8500000" -> "8,500,000.00"). Free-text values for these
# fields pass through unchanged so admin can still type "10% of F08"
# or similar qualifiers when needed.
MONEY_FIELDS: set[str] = {
    "F08",  # Subcontract Price
    "C03",  # Advance Payment Amount
    "C11",  # Rate of Liquidated Damages (AED/day)
    "A07",  # Subcontract Price (mirror)
    "A09",  # Advance Payment (mirror)
    "A20",  # Rate of LDs (mirror)
    # Synthetic tokens (computed at render time as F08 * <pct>/100). Live
    # in MONEY_FIELDS so they pick up the thousands-separator + 2dp
    # formatting like any other AED value.
    "A10_AMOUNT",
    "A21_AMOUNT",
}
# A10 (Performance Security %) and A21 (Maximum LDs %) are percentages,
# NOT money — admin enters "10" / "5" / "15" and the Appendix row computes
# the AED amount via the {{A10_AMOUNT}} / {{A21_AMOUNT}} synthetic tokens
# above.

# Field IDs whose substituted value should render bold so it visually
# matches the hardcoded "M/s. Bhatia General Contracting Co. (L.L.C.)"
# party-name styling in the master. Rev 02 item 3 (Subcontractor row in
# the Appendix) and item 8 (M/s. Microfab in the body) both call for the
# subcontractor company name to be bold wherever it appears.
BOLD_FIELDS: set[str] = {"F02"}

# Field tokens (F01/C03/A15) plus a small set of synthetic tokens like
# {{REFERENCE}} (the agreement's reference stamped into the running header
# on every page). The regex accepts an uppercase letter followed by any mix
# of uppercase letters / digits / underscores so both forms match.
TOKEN_RE = re.compile(r"\{\{\s*([A-Z][A-Z0-9_]*)\s*\}\}")


def _ordinal_suffix(day: int) -> str:
    if 10 <= (day % 100) <= 20:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def _format_money(value: object) -> str:
    """Render a numeric value with thousands separators (and 2-dp decimals).

    Free-form text ("60 days PDC", an empty string, etc.) round-trips
    verbatim — the filter NEVER raises and is safe to chain anywhere.
    Mirrors pdf_service._format_money so the docx and (legacy) WeasyPrint
    pipelines agree on formatting.
    """
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    cleaned = text.replace(",", "").replace(" ", "")
    try:
        n = float(cleaned)
    except ValueError:
        return text
    if n.is_integer():
        return f"{int(n):,}.00"  # Currency context: always show 2dp
    return f"{n:,.2f}"


def _format_longdate(value: object) -> str:
    """Render a date as ``05th May 2026``. Mirrors pdf_service helper."""
    if value is None:
        return ""
    d: date | None = None
    if isinstance(value, datetime):
        d = value
    elif isinstance(value, date):
        d = value
    else:
        text = str(value).strip()
        if not text:
            return ""
        try:
            d = datetime.fromisoformat(text.replace("Z", "+00:00"))
        except ValueError:
            return text
    return f"{d.day:02d}{_ordinal_suffix(d.day)} {d.strftime('%B')} {d.year}"


def _iter_paragraphs(doc) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


_HEADER_FOOTER_ATTRS = (
    "header",
    "first_page_header",
    "even_page_header",
    "footer",
    "first_page_footer",
    "even_page_footer",
)


def _iter_header_footer_text_elements(doc):
    """Yield every ``<w:t>`` element inside every section's headers and
    footers, including those nested inside ``<wps:txbx>`` /
    ``<v:textbox>`` shapes that python-docx doesn't surface through the
    ``header.paragraphs`` API. Required so ``{{REFERENCE}}`` placed in the
    running header stamp gets substituted at render time.
    """
    seen: set[int] = set()
    for section in doc.sections:
        for attr in _HEADER_FOOTER_ATTRS:
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            if not hasattr(hf, "_element"):
                continue
            root = hf._element
            if id(root) in seen:
                continue
            seen.add(id(root))
            for t_el in root.iter(qn("w:t")):
                yield t_el


def _replace_one_value(match: re.Match, values: dict[str, str]) -> str:
    field_id = match.group(1)
    raw = values.get(field_id, "")
    if not raw:
        return ""
    if field_id in DATE_FIELDS:
        return _format_longdate(raw)
    if field_id in MONEY_FIELDS:
        return _format_money(raw)
    return str(raw)


# OOXML text-bearing elements: <w:t> (normal run text) and <w:delText>
# (run text inside a <w:del> track-change). We substitute tokens in both
# so revisions and the surrounding prose stay consistent.
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_TEXT_TAGS = (f"{_W_NS}t", f"{_W_NS}delText")


_LONGDATE_SPLIT_RE = re.compile(r"^(\d{1,2})(st|nd|rd|th)(\s.*)$")


def _split_long_date(value: str) -> tuple[str, str, str] | None:
    """Split a long-form date like ``"05th May 2026"`` into
    ``("05", "th", " May 2026")`` so the suffix can be rendered as a
    superscript run. Returns ``None`` for values that don't match the
    expected ``DD<suffix> Month YYYY`` shape (e.g. an empty value, or
    a date the user typed in a non-canonical format).
    """
    m = _LONGDATE_SPLIT_RE.match(value)
    if not m:
        return None
    return m.group(1), m.group(2), m.group(3)


def _clone_rpr(rpr):
    """Deep-clone a ``<w:rPr>`` element so each emitted run carries its
    own copy of the run-properties tree (font, size, color, italic etc.).
    Returns ``None`` if the source was ``None``.
    """
    if rpr is None:
        return None
    from copy import deepcopy
    return deepcopy(rpr)


def _add_rpr_style(rpr, *, bold: bool = False, superscript: bool = False, highlight: str | None = None):
    """Layer bold / superscript / highlight flags onto an existing ``<w:rPr>``.

    Removes any existing ``<w:b>`` / ``<w:vertAlign>`` / ``<w:highlight>``
    element first so we don't end up with duplicates if the cloned source
    already had them. Creates a fresh ``<w:rPr>`` if none was provided.
    """
    if rpr is None:
        rpr = OxmlElement("w:rPr")
    if bold:
        for existing in rpr.findall(qn("w:b")):
            rpr.remove(existing)
        rpr.append(OxmlElement("w:b"))
    if superscript:
        for existing in rpr.findall(qn("w:vertAlign")):
            rpr.remove(existing)
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    if highlight:
        for existing in rpr.findall(qn("w:highlight")):
            rpr.remove(existing)
        h = OxmlElement("w:highlight")
        h.set(qn("w:val"), highlight)
        rpr.append(h)
    return rpr


def _append_text_node(run_el, text: str) -> None:
    """Append text to a run, converting embedded ``"\\n"`` into ``<w:br/>``
    soft line breaks so multi-line A05/A06 values render across multiple
    visible lines without breaking the surrounding paragraph layout.
    """
    if not text:
        return
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            run_el.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        run_el.append(t)


def _emit_styled_run(
    p_el, base_rpr, text: str, *, bold: bool = False, superscript: bool = False, highlight: str | None = None
):
    """Append a fresh ``<w:r>`` to `p_el` with `text`. Clones `base_rpr`
    so the new run inherits font/size/color from the paragraph's anchor
    run, then layers bold/superscript/highlight on top as requested.
    """
    new_r = OxmlElement("w:r")
    rpr = _clone_rpr(base_rpr)
    if bold or superscript or highlight:
        rpr = _add_rpr_style(rpr, bold=bold, superscript=superscript, highlight=highlight)
    if rpr is not None and len(rpr) > 0:
        new_r.append(rpr)
    _append_text_node(new_r, text)
    p_el.append(new_r)


def _split_cell_soft_breaks(doc: Document) -> None:
    """Expand <w:br/> soft-break sequences in table cell paragraphs into
    separate <w:p> elements.

    After token substitution, multi-line values (C15 bullet lists, A04 project
    details) land in a single <w:p> with <w:br/> between lines. Converting
    each line to its own paragraph gives correct per-line alignment in
    LibreOffice — especially for bullet-style content where wrapped lines must
    start at the same column as the bullet character.
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _split_paras_in_cell(cell)


def _allow_page_breaks_in_ppr(pPr) -> None:
    """Remove keepLines / keepNext / pageBreakBefore flags from a pPr element
    so long multi-paragraph cells (e.g. C15 with many bullet points) are allowed
    to break across pages in LibreOffice. Without this, all generated paragraphs
    inherit the style's keepNext and the entire cell is forced onto one page."""
    for tag in (qn("w:keepLines"), qn("w:keepNext"), qn("w:pageBreakBefore")):
        for el in pPr.findall(tag):
            pPr.remove(el)


def _split_paras_in_cell(cell) -> None:
    from copy import deepcopy
    tc = cell._tc
    for para in list(cell.paragraphs):
        p_el = para._element
        # Quick check: any soft <w:br/> (no w:type attribute) anywhere?
        if not any(
            br.get(qn("w:type")) is None
            for br in p_el.iter(qn("w:br"))
        ):
            continue

        pPr_orig = p_el.find(qn("w:pPr"))

        # Walk paragraph children and segment at soft <w:br/> boundaries.
        # Each segment = list of OOXML elements that form one <w:p>.
        segments: list[list] = [[]]

        for child in list(p_el):
            if child.tag == qn("w:pPr"):
                continue
            if child.tag == qn("w:r"):
                rPr = child.find(qn("w:rPr"))
                cur: list = []  # text/other nodes accumulated for this segment
                for sub in child:
                    if sub.tag == qn("w:rPr"):
                        continue
                    if sub.tag == qn("w:br") and sub.get(qn("w:type")) is None:
                        # Flush accumulated content into a run for the current segment
                        if cur:
                            new_r = OxmlElement("w:r")
                            if rPr is not None:
                                new_r.append(deepcopy(rPr))
                            for t in cur:
                                new_r.append(deepcopy(t))
                            segments[-1].append(new_r)
                            cur = []
                        segments.append([])
                    else:
                        cur.append(sub)
                if cur:
                    new_r = OxmlElement("w:r")
                    if rPr is not None:
                        new_r.append(deepcopy(rPr))
                    for t in cur:
                        new_r.append(deepcopy(t))
                    segments[-1].append(new_r)
            else:
                segments[-1].append(deepcopy(child))

        if len(segments) <= 1:
            continue

        # Insert one new <w:p> per segment before the original paragraph
        insert_pos = list(tc).index(p_el)
        for i, seg_els in enumerate(segments):
            new_p = OxmlElement("w:p")
            if pPr_orig is not None:
                new_pPr = deepcopy(pPr_orig)
                if i > 0:
                    # Remove top-spacing on continuation lines so bullets
                    # don't float apart from one another
                    spc = new_pPr.find(qn("w:spacing"))
                    if spc is not None and qn("w:before") in spc.attrib:
                        del spc.attrib[qn("w:before")]
                # Always allow page breaks in multi-paragraph cells so long
                # content (e.g. C15 with many bullet points) doesn't overflow
                _allow_page_breaks_in_ppr(new_pPr)
                new_p.append(new_pPr)
            for el in seg_els:
                new_p.append(el)
            tc.insert(insert_pos + i, new_p)

        tc.remove(p_el)


def _set_run_text_with_breaks(run, text: str) -> None:
    """Write `text` into `run`, converting "\\n" to OOXML ``<w:br/>`` breaks.

    Used so multifield values (A05/A06 communications addresses) entered as
    multi-line text appear in the PDF as separate lines instead of being
    collapsed to a single line. ``<w:br/>`` is a soft line break — it stays
    inside the host paragraph, preserving table-cell layout.
    """
    r_el = run._r
    for child in list(r_el):
        if child.tag in (qn("w:t"), qn("w:br")):
            r_el.remove(child)
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            r_el.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r_el.append(t)


def _substitute_in_paragraph(
    para: Paragraph, values: dict[str, str], *, highlight_admin_content: bool = False
) -> bool:
    """Replace every ``{{FIELD_ID}}`` token in `para` with its value.

    Pass 1 rebuilds the paragraph's top-level run sequence from styled
    segments so per-field styling can land on substituted values without
    spilling onto the surrounding prose:

    * ``BOLD_FIELDS`` (F02 — the Subcontractor company name) emit a
      dedicated bold run, so M/s. Microfab renders bold inline alongside
      its non-bold surrounding text (Rev 02 items 3 + 8).
    * ``DATE_FIELDS`` (F01) emit three runs: day digits / ordinal
      suffix as superscript / rest of the date, so "05th May 2026"
      renders as "05ᵗʰ May 2026" (Rev 02 item 8).
    * When ``highlight_admin_content`` is set (GM Portal "View PDF", req
      6.3), every substituted field value — every token, since a token is
      admin-entered content by construction — gets a red highlighter
      background (text stays black) so it's visually distinct from the
      surrounding boilerplate. The boilerplate segments themselves are
      never highlighted. Default off: the standard PDF sent to the
      subcontractor/archived must never change.

    Run 0's existing ``<w:rPr>`` is used as the base formatting for every
    emitted run so font / size / color / paragraph-anchored bold all
    carry through.

    Pass 2 then walks every ``<w:t>`` / ``<w:delText>`` text node anywhere
    in the paragraph tree — including text inside ``<w:del>`` / ``<w:ins>``
    track-change wrappers that Phase 4 v2.2 inserts — and substitutes any
    leftover tokens individually so revision spans render with concrete
    values too. Known limitation: Pass 2 mutates text in place with no run
    isolation, so highlighting doesn't reach values inside pending
    track-change spans — acceptable edge case, not blocking.
    """
    p_el = para._element
    changed = False

    runs = para.runs
    if runs:
        full = "".join(r.text or "" for r in runs)
        if "{{" in full:
            matches = list(TOKEN_RE.finditer(full))
            if matches:
                base_rpr = _clone_rpr(runs[0]._r.find(qn("w:rPr")))

                # Clear existing top-level runs (track-change wrappers
                # <w:del>/<w:ins> are not in para.runs, so they survive
                # untouched for Pass 2 to process).
                for r in list(runs):
                    p_el.remove(r._r)

                # Emit segments left-to-right, anchoring each new run on
                # the cloned base run-properties. Only substituted field
                # values get the highlight background — boilerplate
                # segments (the text between tokens) never do.
                value_highlight = "red" if highlight_admin_content else None
                pos = 0
                for m in matches:
                    if m.start() > pos:
                        _emit_styled_run(p_el, base_rpr, full[pos : m.start()])
                    field_id = m.group(1)
                    value = _replace_one_value(m, values)
                    bold = field_id in BOLD_FIELDS
                    if field_id in DATE_FIELDS and value:
                        parts = _split_long_date(value)
                        if parts is not None:
                            day, suffix, rest = parts
                            _emit_styled_run(p_el, base_rpr, day, bold=bold, highlight=value_highlight)
                            _emit_styled_run(
                                p_el, base_rpr, suffix, bold=bold, superscript=True, highlight=value_highlight
                            )
                            _emit_styled_run(p_el, base_rpr, rest, bold=bold, highlight=value_highlight)
                        else:
                            _emit_styled_run(p_el, base_rpr, value, bold=bold, highlight=value_highlight)
                    else:
                        _emit_styled_run(p_el, base_rpr, value, bold=bold, highlight=value_highlight)
                    pos = m.end()
                if pos < len(full):
                    _emit_styled_run(p_el, base_rpr, full[pos:])

                changed = True

    # Pass 2: any token still present inside a track-change wrapper, or
    # inside a nested table cell run we missed, gets substituted on the
    # raw text node so tokens inside <w:del>/<w:ins> resolve too. Newlines
    # inside <w:del>/<w:ins> spans render as spaces — track-change wrappers
    # are not the place for multi-line addresses, so a flat substitution is
    # fine here.
    for tag in _TEXT_TAGS:
        for t_el in p_el.iter(tag):
            txt = t_el.text or ""
            if "{{" not in txt:
                continue
            new_txt = TOKEN_RE.sub(lambda m: _replace_one_value(m, values), txt)
            if new_txt != txt:
                t_el.text = new_txt.replace("\n", " ")
                changed = True

    return changed


def render_agreement_docx_to_pdf(
    values: dict[str, str],
    output_dir: Path | str,
    *,
    master_path: Path | None = None,
    libreoffice_bin: str = "libreoffice",
    timeout_seconds: int = 90,
    accepted_revisions: list[tuple[str, str]] | None = None,
    pending_revisions: list[tuple[str, str, str]] | None = None,
    highlight_admin_content: bool = False,
) -> bytes:
    """Render the SCA PDF by substituting tokens in the master docx and
    converting via LibreOffice headless.

    Parameters
    ----------
    values:
        ``{field_id: entered_value}`` for the agreement. Missing keys cause
        the corresponding token to render empty.
    output_dir:
        Working directory for the intermediate docx + LibreOffice output.
        Caller is responsible for cleanup.
    master_path:
        Override for the master docx location (defaults to the bundled
        ``backend/masters/sca_master_v1.docx``).
    highlight_admin_content:
        GM Portal "View PDF" (req 6.3) — when True, every substituted
        {{FIELD_ID}} value renders in red so admin-entered content is
        visually distinct from boilerplate. Default False — the standard
        PDF (subcontractor-facing, archived) must never change.
    accepted_revisions:
        Phase 4 v2.0 — list of ``(clause_hash, modified_text)`` pairs that
        replace the matching paragraphs in the master before token
        substitution. The modified_text is then itself put through token
        substitution, so it can reference fields the master uses
        (``{{F02}}``, etc.).
    pending_revisions:
        Phase 4 v2.2 — list of ``(clause_hash, original_text, modified_text)``
        tuples. Each target paragraph is wrapped with OOXML
        ``<w:del>``/``<w:ins>`` track-change markers; LibreOffice renders
        them inline as strikethrough + underline. Tokens inside the
        original/modified text are substituted by the iter-based pass in
        ``_substitute_in_paragraph`` so the track-change spans show
        concrete values.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    source = master_path or MASTER_DOCX
    if not source.exists():
        raise FileNotFoundError(f"Master docx not found at {source}")

    doc = Document(str(source))

    # 1) Apply accepted clause revisions BEFORE token substitution so the
    # modified text can contain {{FIELD_ID}} tokens and they'll resolve.
    if accepted_revisions or pending_revisions:
        # Local import to keep the docx_pdf_service module free of the
        # revisions-service dependency unless caller actually passes
        # revisions in.
        from services.clause_revision_service import (
            apply_accepted_revisions_to_doc,
            apply_pending_revisions_to_doc,
        )

        if accepted_revisions:
            apply_accepted_revisions_to_doc(doc, accepted_revisions)

        # 2) Pending revisions get wrapped in <w:del>/<w:ins>. Done AFTER
        # accepted so an accepted revision on the same paragraph wins
        # (which the data layer makes impossible anyway — at most one
        # pending per clause — but we're defensive).
        if pending_revisions:
            apply_pending_revisions_to_doc(doc, pending_revisions)

    # 3) Token substitution — iter-based so it descends into the
    # <w:del>/<w:ins> blocks just inserted and resolves tokens in the
    # revision text spans too.
    for para in _iter_paragraphs(doc):
        _substitute_in_paragraph(para, values, highlight_admin_content=highlight_admin_content)

    # 3.5) Expand any <w:br/> soft-break sequences that token substitution
    # produced in table cells (e.g. C15 bullets, A04 project details) into
    # separate <w:p> elements so each line aligns properly in LibreOffice.
    _split_cell_soft_breaks(doc)

    # 4) Headers / footers (including textbox content). The {{REFERENCE}}
    # token stamped on every page lives inside <wps:txbx> shapes that the
    # body-paragraph walker doesn't reach. Substitute directly on each
    # <w:t> element under every section's header/footer.
    for t_el in _iter_header_footer_text_elements(doc):
        txt = t_el.text or ""
        if "{{" not in txt:
            continue
        new_txt = TOKEN_RE.sub(lambda m: _replace_one_value(m, values), txt)
        if new_txt != txt:
            # Headers/footers are single-line — drop any embedded newlines
            # that a value (e.g. an A05 address) might accidentally carry.
            t_el.text = new_txt.replace("\n", " ")

    intermediate_docx = output_dir / "rendered.docx"
    doc.save(str(intermediate_docx))

    # LibreOffice insists on a writable "user installation" profile dir.
    # Without -env:UserInstallation it tries to create one under $HOME,
    # which fails when the process runs as a service user with /var/www
    # as its home.
    lo_profile = output_dir / "_lo_profile"
    lo_profile.mkdir(exist_ok=True)
    cmd = [
        libreoffice_bin,
        f"-env:UserInstallation=file://{lo_profile}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(intermediate_docx),
    ]
    proc = subprocess.run(cmd, capture_output=True, timeout=timeout_seconds)
    if proc.returncode != 0:
        raise RuntimeError(
            "LibreOffice conversion failed (rc=%d): %s"
            % (proc.returncode, proc.stderr.decode("utf-8", errors="replace"))
        )

    pdf_path = intermediate_docx.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError(
            f"LibreOffice reported success but no PDF at {pdf_path}; "
            f"stdout: {proc.stdout.decode('utf-8', errors='replace')[:500]}"
        )
    return pdf_path.read_bytes()
