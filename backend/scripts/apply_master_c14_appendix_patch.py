#!/usr/bin/env python3
"""Idempotent patch: add {{C14}} to Performance Security appendix row.

The master docx Performance Security cell (Table 3, Row 11, Cell 2) shows:
  {{A10}}% of the Contract Price
  {{A10_AMOUNT}} AED

This patch appends a third line so the security instrument type is visible:
  {{A10}}% of the Contract Price
  {{A10_AMOUNT}} AED
  {{C14}}

Run once; re-running is safe (checks for the sentinel before writing).
"""
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"
BAK_DOCX = MASTER_DOCX.with_suffix(".docx.pre-c14.bak")
SENTINEL = "{{C14}}"


def main() -> None:
    doc = Document(str(MASTER_DOCX))

    # Performance Security row: Table 3, Row 11, Cell 2
    try:
        cell = doc.tables[3].rows[11].cells[2]
    except IndexError:
        raise SystemExit("ERROR: Could not locate Table 3 / Row 11. Master structure may have changed.")

    para = cell.paragraphs[0]

    if SENTINEL in para.text:
        print(f"Already patched — '{SENTINEL}' present in Performance Security cell. Nothing to do.")
        return

    if not para.runs:
        raise SystemExit("ERROR: Performance Security cell has no runs. Master structure unexpected.")

    # Back up before modifying
    if not BAK_DOCX.exists():
        import shutil
        shutil.copy2(str(MASTER_DOCX), str(BAK_DOCX))
        print(f"Backup written to {BAK_DOCX.name}")

    # Append <w:br/> + <w:t>{{C14}}</w:t> to the first (and only) run
    r_el = para.runs[0]._r

    br = OxmlElement("w:br")
    r_el.append(br)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = SENTINEL
    r_el.append(t)

    doc.save(str(MASTER_DOCX))
    print(f"Patched: '{SENTINEL}' added to Performance Security appendix row.")
    print(f"Cell now reads: {repr(para.text)}")


if __name__ == "__main__":
    main()
