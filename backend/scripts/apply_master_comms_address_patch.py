"""
apply_master_comms_address_patch.py

Adds the fixed Main Contractor Address block BEFORE the Subcontractor
Address in the appendix table "Communications Address for Serving of the
Notices" cell (Table 3, row 6, col 2).

Fixed content (not a field):
  The Main Contractor Address
  Attention: Mr.Moath Milhem
  Position Title: General Manager
  Address: Ras Al Khor, Dubai
  Facsimile Number: +9714 333 1027
  Email Address: milhem@bhatiacompany.com
                 info@bhatiacompany.com
  (blank line)
  The Subcontractor Address
  {{A06}}

Idempotent: guards on "Mr.Moath Milhem" already being present.
Creates .pre-comms-addr.bak on first run.
"""

import copy, shutil
from pathlib import Path
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MASTER_PATH = Path("backend/masters/sca_master_v1.docx")
BAK_SUFFIX  = ".pre-comms-addr.bak"

W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag):
    return f"{{{W}}}{tag}"


def _make_rPr(bold=False):
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    sz = OxmlElement("w:sz")
    sz.set(_w("val"), "20")          # 10 pt
    szCs = OxmlElement("w:szCs")
    szCs.set(_w("val"), "20")
    w_elem = OxmlElement("w:w")
    w_elem.set(_w("val"), "105")
    rPr.append(sz)
    rPr.append(szCs)
    rPr.append(w_elem)
    return rPr


def _make_run(text, bold=False, add_br=True):
    """Return a <w:r> element with optional <w:br/> appended."""
    r = OxmlElement("w:r")
    r.append(_make_rPr(bold=bold))
    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    if add_br:
        r.append(OxmlElement("w:br"))
    return r


def build_new_cell_xml(tc):
    """
    Replace the content of the given <w:tc> element with the two-address layout.
    The <w:tcPr> is preserved unchanged.
    """
    # Remove existing paragraphs from the cell
    for p in tc.findall(_w("p")):
        tc.remove(p)

    # Build the replacement paragraph
    p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(_w("val"), "TableParagraph")
    spacing = OxmlElement("w:spacing")
    spacing.set(_w("before"), "71")
    spacing.set(_w("after"), "0")
    pPr.append(pStyle)
    pPr.append(spacing)
    p.append(pPr)

    # ── Main Contractor block ──────────────────────────────────────────────
    p.append(_make_run("The Main Contractor Address", bold=True, add_br=True))
    p.append(_make_run("Attention: Mr.Moath Milhem",  bold=False, add_br=True))
    p.append(_make_run("Position Title: General Manager", bold=False, add_br=True))
    p.append(_make_run("Address: Ras Al Khor, Dubai",     bold=False, add_br=True))
    p.append(_make_run("Facsimile Number: +9714 333 1027", bold=False, add_br=True))
    p.append(_make_run("Email Address: milhem@bhatiacompany.com", bold=False, add_br=True))
    p.append(_make_run("                  info@bhatiacompany.com", bold=False, add_br=True))

    # Blank line separator
    r_blank = OxmlElement("w:r")
    r_blank.append(_make_rPr(bold=False))
    t_blank = OxmlElement("w:t")
    t_blank.text = ""
    r_blank.append(t_blank)
    r_blank.append(OxmlElement("w:br"))
    p.append(r_blank)

    # ── Subcontractor block ────────────────────────────────────────────────
    p.append(_make_run("The Subcontractor Address", bold=True, add_br=True))
    p.append(_make_run("{{A06}}", bold=False, add_br=False))

    tc.append(p)


def main():
    bak = MASTER_PATH.with_suffix(MASTER_PATH.suffix + BAK_SUFFIX)
    if not bak.exists():
        shutil.copy2(MASTER_PATH, bak)
        print(f"Backup created: {bak.name}")
    else:
        print(f"Backup already exists: {bak.name} (skipping)")

    master = Document(str(MASTER_PATH))
    tbl    = master.tables[3]
    cell   = tbl.rows[6].cells[2]

    if "Mr.Moath Milhem" in cell.text:
        print("Main Contractor address already present – nothing to do.")
        return

    if "The Subcontractor Address" not in cell.text:
        print("WARNING: row 6 col 2 doesn't look like expected – check the script.")
        print(f"  Actual text: {repr(cell.text[:120])}")
        return

    print("Patching Communications Address cell …")
    build_new_cell_xml(cell._tc)

    master.save(str(MASTER_PATH))
    print("Saved master.")

    # Quick verify
    master2 = Document(str(MASTER_PATH))
    patched = master2.tables[3].rows[6].cells[2].text
    print(f"Cell text after patch:\n{patched}")


if __name__ == "__main__":
    main()
