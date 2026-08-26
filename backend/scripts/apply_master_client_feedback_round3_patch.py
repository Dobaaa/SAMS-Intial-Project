"""Client feedback round 3 (2026-08-26), 7 docx items + 1 same-day
follow-up (Subcontract Quantities row left-indent). Idempotent — each
sub-fix checks its own precondition before touching anything, matching the
pattern established in apply_master_remove_commencement_time_pdc_patch.py.
Run directly: `python scripts/apply_master_client_feedback_round3_patch.py`.
"""
from __future__ import annotations

import copy
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

MASTER_PATH = Path(__file__).resolve().parent.parent / "masters" / "sca_master_v1.docx"
BACKUP_PATH = MASTER_PATH.with_suffix(".docx.pre-round3.bak")


def _find_table_cell(doc, row_label_substr: str, col: int):
    for table in doc.tables:
        for row in table.rows:
            if row_label_substr in row.cells[0].text:
                return row.cells[col]
    return None


def _fix_pdc_appendix_rows(doc) -> bool:
    """Req 1: appendix Table 3's Interim/1st-Half/2nd-Half PDC rows should
    read "X days PDC within 15 days from the invoice date", not just the
    bare {{C05}}/{{C06}}/{{C07}} number."""
    targets = [
        ("Interim Payment to be paid", "{{C05}}", "{{C05}} days PDC within 15 days from the invoice date"),
        ("1st Half of retention Money", "{{C06}}", "{{C06}} days PDC within 15 days from the invoice date"),
        ("2nd Half of retention Money", "{{C07}}", "{{C07}} days PDC within 15 days from the invoice date"),
    ]
    changed = False
    for label_substr, old_token, new_text in targets:
        cell = _find_table_cell(doc, label_substr, 2)
        if cell is None:
            continue
        for p in cell.paragraphs:
            if p.text.strip() == old_token:
                for r in list(p.runs):
                    r.text = ""
                p.runs[0].text = new_text
                changed = True
    return changed


def _fix_pdc_body_wording(doc) -> bool:
    """Req 1: "{{C05}}-day Post-Dated Cheque" -> "{{C05}} days Post-Dated
    Cheque" (also C06/C07) in the 3.4.6/3.4.7 body clauses."""
    changed = False
    for token in ("C05", "C06", "C07"):
        old = f"{{{{{token}}}}}-day Post-Dated Cheque"
        new = f"{{{{{token}}}}} days Post-Dated Cheque"
        for p in doc.paragraphs:
            if old in p.text:
                full = "".join(r.text for r in p.runs)
                if old not in full:
                    continue
                new_full = full.replace(old, new)
                for r in list(p.runs)[1:]:
                    r.text = ""
                p.runs[0].text = new_full
                changed = True
    return changed


def _fix_comms_address_email_alignment(doc) -> bool:
    """Req 3: the second Main Contractor email ("info@...") was pushed
    right via literal leading spaces, meant to line it up under
    "milhem@..." on the assumption both emails sit on the same visual
    line after "Email Address: ". They don't: at this cell's bold weight
    "Email Address: milhem@bhatiacompany.com" is too wide for the column
    and word-wraps, so "milhem@..." itself starts flush at the paragraph's
    left edge, not indented. True alignment is therefore no indent at
    all — verified by rendering both the before and after state to PDF
    and comparing crops, not by eye on the source XML."""
    cell = _find_table_cell(doc, "Communications Address for Serving", 2)
    if cell is None:
        return False
    for r in cell.paragraphs[0].runs:
        if r.text.lstrip(" ").startswith("info@bhatiacompany.com") and r.text.startswith(" "):
            r.text = r.text.lstrip(" ")
            return True
    return False


def _fix_subcontractor_contact_person(doc) -> bool:
    """Req 4: the Subcontractor Address block has no Contact Person line,
    unlike the Main Contractor block's "Attention:" line. Insert one,
    sourced from the new {{SUBCONTRACTOR_CONTACT_PERSON}} synthetic token
    (wired in pdf_service._build_value_map)."""
    cell = _find_table_cell(doc, "Communications Address for Serving", 2)
    if cell is None:
        return False
    p = cell.paragraphs[0]
    header_run = None
    for r in p.runs:
        if r.text.strip() == "The Subcontractor Address":
            header_run = r
            break
    if header_run is None:
        return False
    if "SUBCONTRACTOR_CONTACT_PERSON" in p.text:
        return False

    # Clone the "Attention:" line's (non-bold) run formatting.
    template_rpr = None
    for r in p.runs:
        if r.text.startswith("Attention:"):
            template_rpr = r._r.find(qn("w:rPr"))
            break

    new_r = OxmlElement("w:r")
    if template_rpr is not None:
        new_r.append(copy.deepcopy(template_rpr))
    t_el = OxmlElement("w:t")
    t_el.text = "Contact Person: {{SUBCONTRACTOR_CONTACT_PERSON}}"
    new_r.append(t_el)
    new_r.append(OxmlElement("w:br"))

    header_run._r.addnext(new_r)
    return True


def _convert_retention_items_to_bullets(doc) -> bool:
    """2026-08-26 follow-up: "Retention Money and Final Payment" (c)'s two
    sub-items (roman numerals I/II — the first/second 50% retention PDC
    paragraphs) should render as bullet points, tabbed clearly under (c)
    as sub-bullets. They're already nested under (c) in the numbering
    (ilvl=4 vs (c)'s ilvl=3, same numId=48), but that shared numId=48/
    ilvl=4 definition is also used elsewhere (Call on Bonds' I/II/III at
    paragraphs 288-290), so its abstract numbering format can't be changed
    without also affecting that unrelated clause. Instead: detach just
    these two paragraphs from the numbered list (numPr removed) and give
    each a literal "•" + a real tab stop, with a left-indent well past
    (c)'s own (c) sits at left=2356/hanging=339 twips, i.e. its own
    letter starts at ~2017; these get left=3560/hanging=460, i.e. bullet
    at ~3100 — clearly further right/"tabbed in" than (c) itself, unlike
    the roman-numeral markers which sat at ~1973-1991, almost exactly
    level with (c)."""
    from docx.oxml.ns import qn as _qn

    target_texts = (
        "The first 50% of the retention shall be released by",
        "The remaining 50% of the retention shall be released by",
    )
    changed = False
    for p in doc.paragraphs:
        if not any(p.text.startswith(t) for t in target_texts):
            continue
        pPr = p._p.pPr
        if pPr is None:
            continue
        numPr = pPr.find(_qn("w:numPr"))
        if numPr is None:
            continue  # already converted (idempotent)

        pPr.remove(numPr)

        ind = pPr.find(_qn("w:ind"))
        if ind is not None:
            ind.set(_qn("w:hanging"), "460")
            ind.set(_qn("w:left"), "3560")

        tabs = pPr.find(_qn("w:tabs"))
        if tabs is not None:
            for tab in tabs.findall(_qn("w:tab")):
                if tab.get(_qn("w:val")) == "left":
                    tab.set(_qn("w:pos"), "3560")

        run0 = p.runs[0]
        r_el = run0._r
        rPr = r_el.find(_qn("w:rPr"))
        insert_at = list(r_el).index(rPr) + 1 if rPr is not None else 0

        bullet_t = OxmlElement("w:t")
        bullet_t.set(_qn("xml:space"), "preserve")
        bullet_t.text = "•"
        tab_el = OxmlElement("w:tab")
        r_el.insert(insert_at, tab_el)
        r_el.insert(insert_at, bullet_t)
        changed = True
    return changed


def _fix_quantities_row_left_indent(doc) -> bool:
    """2026-08-26 follow-up: "The Subcontract Quantities (... or
    Re-measurable)" row's value cell ({{C02}}) has left_indent=3810 EMU,
    while every sibling value cell in this table (The Subcontract Price,
    Advance Payment Amount, Performance Security, ...) uses ~65000 EMU —
    about 17x more. That's why {{C02}}'s value ("Lump Sum"/"Re-measurable")
    sits flush against the cell border while every other row's value has a
    visible left margin ("needs a space before it")."""
    from docx.shared import Emu

    target = Emu(65405)
    cell = _find_table_cell(doc, "The Subcontract Quantities", 2)
    if cell is None:
        return False
    changed = False
    for p in cell.paragraphs:
        if p.paragraph_format.left_indent != target:
            p.paragraph_format.left_indent = target
            changed = True
    return changed


def _fix_dlp_appendix_months(doc) -> bool:
    """Req 5: Appendix Table 4's Defects Liability Period row shows the
    bare number; the body clause already appends "Months" (para ~443).
    Add the unit to the appendix row to match."""
    cell = _find_table_cell(doc, "Defects Liability Period", 2)
    if cell is None:
        return False
    for p in cell.paragraphs:
        if p.text.strip() == "{{A19}}":
            for r in list(p.runs)[1:]:
                r.text = ""
            p.runs[0].text = "{{A19}} months"
            return True
    return False


def _trim_blank_run_after(doc, anchor_prefix: str, max_blanks_expected: int = 1) -> bool:
    """Find the paragraph starting with `anchor_prefix`, then remove extra
    blank paragraphs immediately following it down to `max_blanks_expected`.
    Anchor-based (not a generic scan) so it can't touch an unrelated
    double-blank elsewhere in the 40+ page document."""
    paras = doc.paragraphs
    anchor_idx = None
    for i, p in enumerate(paras):
        if p.text.startswith(anchor_prefix):
            anchor_idx = i
            break
    if anchor_idx is None:
        return False

    blank_idxs = []
    j = anchor_idx + 1
    while j < len(paras) and paras[j].text.strip() == "":
        blank_idxs.append(j)
        j += 1

    if len(blank_idxs) <= max_blanks_expected:
        return False

    for idx in blank_idxs[max_blanks_expected:]:
        el = paras[idx]._p
        el.getparent().remove(el)
    return True


def _trim_advance_payment_double_blanks(doc) -> bool:
    """Req 6: two stray double-blank-paragraph gaps in 3.4.1 (Advance
    Payment) — after item (d) and after item (f) — each should be a single
    blank, matching every other list section's spacing."""
    changed_d = _trim_blank_run_after(doc, "Advance Payment Guarantee shall be in the form")
    changed_f = _trim_blank_run_after(doc, "The Subcontractor acknowledges that the advance payment remains the property")
    return changed_d or changed_f


def main() -> None:
    if not MASTER_PATH.exists():
        raise SystemExit(f"Master docx not found: {MASTER_PATH}")

    doc = Document(str(MASTER_PATH))

    results = {
        "pdc_appendix_rows": _fix_pdc_appendix_rows(doc),
        "pdc_body_wording": _fix_pdc_body_wording(doc),
        "comms_email_alignment": _fix_comms_address_email_alignment(doc),
        "subcontractor_contact_person": _fix_subcontractor_contact_person(doc),
        "dlp_appendix_months": _fix_dlp_appendix_months(doc),
        "advance_payment_blanks": _trim_advance_payment_double_blanks(doc),
        "quantities_row_left_indent": _fix_quantities_row_left_indent(doc),
        "retention_items_bullets": _convert_retention_items_to_bullets(doc),
    }

    any_changed = any(results.values())
    print("Sub-fix results:", results)

    if not any_changed:
        print("No changes needed (already applied).")
        return

    if not BACKUP_PATH.exists():
        BACKUP_PATH.write_bytes(MASTER_PATH.read_bytes())
        print(f"Backup written: {BACKUP_PATH}")

    doc.save(str(MASTER_PATH))
    print(f"Saved: {MASTER_PATH}")

    # ---- verification ----
    # doc.paragraphs only walks direct body-level <w:p> elements — it does
    # NOT descend into table cells, so checks on table content must go
    # through doc.tables instead (same lesson as the SDT/TOC bug logged
    # 2026-08-24: never trust doc.paragraphs alone for verification).
    doc2 = Document(str(MASTER_PATH))
    body_text = "\n".join(p.text for p in doc2.paragraphs)
    table_text = "\n".join(
        cell.text for table in doc2.tables for row in table.rows for cell in row.cells
    )
    assert "{{C05}}-day" not in body_text, "stale hyphenated C05 wording"
    assert "{{C06}}-day" not in body_text, "stale hyphenated C06 wording"
    assert "{{C07}}-day" not in body_text, "stale hyphenated C07 wording"
    assert "{{C05}} days Post-Dated Cheque" in body_text
    assert "SUBCONTRACTOR_CONTACT_PERSON" in table_text
    assert "•The first 50% of the retention" in body_text.replace("•\t", "•"), "retention item I not bulleted"
    assert "•The remaining 50% of the retention" in body_text.replace("•\t", "•"), "retention item II not bulleted"

    for table in doc2.tables:
        for row in table.rows:
            if "Interim Payment to be paid" in row.cells[0].text:
                assert "days PDC within 15 days from the invoice date" in row.cells[2].text
            if "Defects Liability Period" in row.cells[0].text and "{{A19}}" in row.cells[2].text:
                assert "months" in row.cells[2].text
            if "The Subcontract Quantities" in row.cells[0].text:
                from docx.shared import Emu
                for p in row.cells[2].paragraphs:
                    assert p.paragraph_format.left_indent == Emu(65405), "quantities row indent not fixed"
    print("Verification passed.")


if __name__ == "__main__":
    main()
