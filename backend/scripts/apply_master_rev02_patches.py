"""Apply Rev 02 corrections to the SCA master docx.

Rev 02 round of BGCC feedback (18.05.2026) flagged four defects in the
rendered PDF that all live in the master template, not in the rendering
code. This script applies the surgical fixes in one reproducible pass:

* **Item 1** — Communications Address (clause 1.6) rows in the Appendix
  used literal "…………………………" dots, so admin's A05 / A06 multifield
  inputs never reached the PDF. We collapse the 5 sub-rows per side into
  one row per side and place ``{{A05}}`` / ``{{A06}}`` tokens. The
  service's substitution layer converts the multi-line textarea value
  into ``<w:br/>`` soft line breaks so each line still renders distinct.

* **Item 2** — "Time for Completion" and "Sections (Milestones)" rows in
  the Appendix used literal dots and "MS1 / MS2" placeholders. We replace
  the cell text to interpolate ``{{A16}}`` (Project), ``{{A17}}``
  (Subcontract Works) and ``{{A18}}`` (Milestones).

* **Performance Security %** + **Maximum LDs %** — A10 (Performance
  Security) and A21 (Maximum LDs) are percentage inputs, but their
  Appendix rows used to hardcode ``"10% of the Contract Price / … AED"``.
  Rewrite both rows to interpolate the entered percentage **and** an
  AED amount derived from F08 via the new ``{{A10_AMOUNT}}`` /
  ``{{A21_AMOUNT}}`` synthetic tokens (computed in
  ``services.pdf_service._inject_percentage_amounts`` and rendered as
  money). Final cell text reads e.g. ``"10% of the Contract Price\n
  850,000.00 AED"`` for F08=8.5M with A10=10.

* **Item 11** — Sub-clause 3.4(e) "Retention Money and Final Payment"
  was meant to start on a fresh page, but the client reverted the
  request after seeing the rendered result. The helper that inserts
  the break is kept in this module for future re-enablement, but is
  no longer called from ``main()``.

* **Blank section-title page** — between the cover and the actual
  Form content the master had a dedicated section holding only a
  "Form of Subcontract Agreement" title and the BGCC subtitle, which
  rendered as a visually empty page 2. Delete that section's
  paragraphs (and its terminating sectPr) so the form content's
  first page becomes page 2 in the rendered PDF.

* **Cover page layout** — the three label paragraphs
  ("The Project", "Subcontractor Name", "Scope Title") used a
  tab-stop between label and colon, and when the value wrapped, the
  continuation line fell back to the left margin (looked broken).
  Replace the three paragraphs with a 3-row borderless 2-column
  table — label cell ~2 inches wide and bold, value cell takes the
  remaining width — so each entry stays a coherent row and wrapped
  values indent under the value column.

* **Header stamp** — every page's running header reads
  "BGCC P-XXX / SCA#-ZZZ" as a static placeholder. Replace it with a
  ``{{REFERENCE}}`` token so the rendered PDF stamps the agreement's
  reference number (e.g. ``SAG-2026-319-001``) on every page. The
  "Page N of 41" total next to it is intentionally left alone.

Run::

    python backend/scripts/apply_master_rev02_patches.py
"""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"
BACKUP_DOCX = MASTER_DOCX.with_suffix(".docx.pre-rev02.bak")

# Runs in the "BGCC P-XXX / SCA#-ZZZ" stamp, in order. The first non-space
# run is rewritten to {{REFERENCE}}; the rest are cleared. Whitespace runs
# stay so the row's tab + page-of indicator on the right keeps its layout.
_STAMP_TOKENS = ("BGCC", "P-XXX", "/", "SCA#-", "ZZZ")


def _set_cell_text(cell, text: str) -> None:
    """Replace cell content with a single paragraph holding `text`.

    Preserves the first paragraph's existing run-formatting (so the
    Tahoma face + size + alignment of the original cell stays intact)
    and rewrites the text; clears any other paragraphs.
    """
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _remove_row(table, row) -> None:
    table._tbl.remove(row._tr)


def _insert_page_break_before(paragraph) -> None:
    """Insert a paragraph holding only a ``<w:br w:type="page"/>`` before `paragraph`."""
    new_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    new_p.append(r)
    paragraph._element.addprevious(new_p)


def patch_appendix(doc) -> None:
    """Items 1 & 2: rewire Appendix Table 2 rows to use tokens."""
    appendix = doc.tables[2]  # The 28-row "Item Description | Clause | Info" table.
    rows = list(appendix.rows)

    # Item 1 — Main Contractor address: R6 keeps header + {{A05}}; delete R7-R10.
    _set_cell_text(rows[6].cells[2], "The Main Contractor Address\n{{A05}}")
    for r in (rows[10], rows[9], rows[8], rows[7]):  # remove back-to-front
        _remove_row(appendix, r)

    # After the four deletions, R11..R16 shift to indices 7..12. Re-fetch.
    rows = list(appendix.rows)

    # Item 1 — Subcontractor address: row that was R11 (now R7) keeps header +
    # {{A06}}; delete the five sub-rows that followed.
    _set_cell_text(rows[7].cells[2], "The Subcontractor Address\n{{A06}}")
    for r in (rows[12], rows[11], rows[10], rows[9], rows[8]):
        _remove_row(appendix, r)

    # Item 2 — Time for Completion + Milestones. Locate by label so the lookup
    # still works after the deletions above shifted indices.
    # Performance Security — also locate by label here. The row's data cell
    # held only static "10% of the Contract Price / … AED" text; rewrite to
    # interpolate {{A10}} (the cascaded AED amount).
    rows = list(appendix.rows)
    for row in rows:
        label = row.cells[0].text.strip().lower()
        if label.startswith("time for completion") and "section" not in label:
            _set_cell_text(
                row.cells[2],
                "Time for Completion of the Project: {{A16}}\n"
                "Time for Completion of the Subcontract Works: {{A17}}",
            )
        elif "milestones" in label or "sections" in label:
            _set_cell_text(row.cells[2], "{{A18}}")
        elif label == "performance security":
            _set_cell_text(
                row.cells[2],
                "{{A10}}% of the Contract Price\n{{A10_AMOUNT}} AED",
            )

    # Maximum Liquidated Damages lives in Table 3 (the "Defects Liability /
    # LDs / Insurance / Dispute" table), not Table 2. Same percentage-+-amount
    # treatment as Performance Security.
    if len(doc.tables) > 3:
        max_lds_table = doc.tables[3]
        for row in max_lds_table.rows:
            label = row.cells[0].text.strip().lower()
            if label == "maximum liquidated damages":
                _set_cell_text(
                    row.cells[2],
                    "{{A21}}% of the Contract Price\n{{A21_AMOUNT}} AED",
                )


_COVER_ROWS = (
    ("The Project:", "{{F06}}"),
    ("Subcontractor Name:", "{{F02}}"),
    ("Scope Title:", "{{F09}}"),
)


# Cover-page run typography (matches the original master's three cover-
# label paragraphs verbatim — Times New Roman 11pt, white text overlaid on
# the cover's branded backdrop). Kept verbatim so the new table-based
# layout is visually indistinguishable from the prior paragraph layout
# apart from how wrapped values flow.
_COVER_RUN_FONT = "Times New Roman"
_COVER_RUN_COLOR = "FFFFFF"
_COVER_RUN_SZ = "22"  # half-points; 22 = 11pt

# Table column widths, in twentieths-of-a-point (twips, 1440 = 1 inch).
# Label column ~1.9" + value column ~3.1" = ~5 inches total. The original
# paragraph constrained text via w:ind right=4358 (≈3.3" usable width);
# we keep the table within that left-side column, just a hair wider so
# common values like the subcontractor company name fit on one line.
_COVER_LABEL_W = "2736"  # 1.9"
_COVER_VALUE_W = "4464"  # 3.1"
_COVER_TABLE_W = "7200"  # sum (5")
_COVER_TABLE_INDENT = "262"  # mirrors the original w:ind left=262


def _make_cover_table_xml():
    """Build a 3-row, 2-column borderless table for the cover-page entries.

    Each row holds (label) | (value with token). Cells inherit the
    original cover styling (Times New Roman 11pt, white text) so they
    sit cleanly over the existing cover backdrop. The whole table is
    narrower than the body width — only ~4" — to stay inside the
    cover's left-side column where the original labels lived.
    """
    tbl = OxmlElement("w:tbl")

    tbl_pr = OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), _COVER_TABLE_W)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    # Indent the table so its left edge sits where the original
    # paragraphs sat (w:ind left=262 in the master).
    tbl_indent = OxmlElement("w:tblInd")
    tbl_indent.set(qn("w:w"), _COVER_TABLE_INDENT)
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_indent)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    tbl.append(tbl_pr)

    tbl_grid = OxmlElement("w:tblGrid")
    for w in (_COVER_LABEL_W, _COVER_VALUE_W):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), w)
        tbl_grid.append(gc)
    tbl.append(tbl_grid)

    for label, value in _COVER_ROWS:
        tr = OxmlElement("w:tr")
        for cell_text, width in (
            (label, _COVER_LABEL_W),
            (value, _COVER_VALUE_W),
        ):
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), width)
            tc_w.set(qn("w:type"), "dxa")
            tc_pr.append(tc_w)
            tc.append(tc_pr)

            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), _COVER_RUN_FONT)
            r_fonts.set(qn("w:hAnsi"), _COVER_RUN_FONT)
            r_pr.append(r_fonts)
            color = OxmlElement("w:color")
            color.set(qn("w:val"), _COVER_RUN_COLOR)
            r_pr.append(color)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), _COVER_RUN_SZ)
            r_pr.append(sz)
            r.append(r_pr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    return tbl


def patch_cover_page_layout(doc) -> None:
    """Rev 02 styling fix: replace the three cover-page label paragraphs
    with a borderless 2-column table so each entry stays one logical
    row and wrapped values indent under the value column instead of
    falling back to the left margin.
    """
    targets: list = []
    seen = set()
    for p in doc.paragraphs:
        text = p.text.strip()
        for label in ("The Project", "Subcontractor Name", "Scope Title"):
            if text.startswith(label) and label not in seen:
                targets.append(p)
                seen.add(label)
                break
        if len(seen) == 3:
            break
    if len(targets) != 3:
        raise RuntimeError(
            "Could not locate all 3 cover-page label paragraphs "
            f"(found {len(targets)})."
        )

    parent = targets[0]._element.getparent()
    insert_at = list(parent).index(targets[0]._element)
    parent.insert(insert_at, _make_cover_table_xml())

    for p in targets:
        parent.remove(p._element)


def patch_remove_form_title_page(doc) -> None:
    """Delete the "Form of Subcontract Agreement" section-title page.

    The master's structure between the cover and the actual Form
    content is:

        ...Annexures...                       (last cover content)
        <sectPr nextPage>                     (ends Section 1 = cover)
        Bhatia General Contracting Co. L.L.C. (BGCC     ← P20
        (10 blank paragraphs)
        Form of Subcontract Agreement
        (10 blank paragraphs)
        {{REFERENCE}}\\tPage 1 of 41 <sectPr nextPage>  ← P42
        (2 blank paragraphs)
        FORM OF SUBCONTRACT AGREEMENT THIS SUBCONTRACT…  ← real form

    The whole Section 2 (P20 through P42 inclusive) is visually a
    near-blank divider page. We remove every paragraph in that range,
    including the sectPr-bearing P42, so what was Section 3 (form
    body) inherits page numbering directly from Section 1's terminator
    and renders as page 2 of the PDF.
    """
    paragraphs = doc.paragraphs

    # Anchor: the paragraph that starts Section 2 content.
    start_idx = next(
        (
            i for i, p in enumerate(paragraphs)
            if p.text.strip().startswith("Bhatia General Contracting Co.")
        ),
        None,
    )
    if start_idx is None:
        raise RuntimeError(
            "Could not find Section 2 anchor 'Bhatia General Contracting Co.' "
            "to remove the section-title page."
        )

    # End: the first paragraph at or after start_idx that carries a sectPr.
    end_idx = None
    for i in range(start_idx, len(paragraphs)):
        sectPr = paragraphs[i]._element.find(".//" + qn("w:sectPr"))
        if sectPr is not None:
            end_idx = i
            break
    if end_idx is None:
        raise RuntimeError(
            "Could not find terminating sectPr for the section-title page."
        )

    body = paragraphs[start_idx]._element.getparent()
    for p in paragraphs[start_idx : end_idx + 1]:
        body.remove(p._element)


def patch_clause_3_4_e_pagebreak(doc) -> None:
    """Item 11: page-break before "Retention Money and Final Payment"."""
    for p in doc.paragraphs:
        if p.text.strip().lower().startswith("retention money and final payment"):
            _insert_page_break_before(p)
            return
    raise RuntimeError(
        "Could not find 'Retention Money and Final Payment' heading to "
        "insert the 3.4(e) page break before."
    )


def _rewrite_stamp_runs_in_paragraph(p) -> bool:
    """If a paragraph holds the BGCC P-XXX / SCA#-ZZZ stamp split across
    runs, collapse the substantive tokens into a single ``{{REFERENCE}}``
    run and clear the others. Returns True if a change was made.
    """
    runs = p.runs
    indices = [
        i for i, r in enumerate(runs)
        if (r.text or "").strip() in _STAMP_TOKENS
    ]
    # All five stamp tokens (or at least 4 — be tolerant of small variants)
    # must be present, in order, for this to be the stamp paragraph.
    if len(indices) < 4:
        return False
    first = indices[0]
    runs[first].text = "{{REFERENCE}}"
    for idx in indices[1:]:
        runs[idx].text = ""
    return True


def patch_body_stamps(doc) -> None:
    """Pages 1 and 4 carry the running header inline in the body
    ("BGCC P-XXX / SCA#-ZZZ\\tPage N of 41") instead of in a footer xml.
    Replace the BGCC stamp segment with a {{REFERENCE}} token, leaving
    the "Page N of 41" portion intact.
    """
    for p in doc.paragraphs:
        if "BGCC" in p.text and "P-XXX" in p.text:
            _rewrite_stamp_runs_in_paragraph(p)


# Pattern matches consecutive <w:t> runs holding the stamp tokens, including
# the small spacing/punctuation runs between them. We rewrite the first run
# to {{REFERENCE}} and erase the rest in one substitution.
_FOOTER_STAMP_RE = re.compile(
    r"(<w:t[^>]*>)BGCC(</w:t>)"            # opening run with "BGCC"
    r"(?P<between>.*?)"                     # everything up to ZZZ (greedy-min)
    r"(<w:t[^>]*>)ZZZ(</w:t>)",
    re.DOTALL,
)


def _rewrite_stamp_in_xml(xml_text: str) -> str:
    """Replace each ``BGCC ... ZZZ`` stamp occurrence with a single run
    containing ``{{REFERENCE}}`` and emptied trailing runs. We keep the
    surrounding ``<w:r>``/``<w:rPr>`` structure for the first run intact
    so the bold + size styling of the stamp carries through to the
    substituted reference at render time. Empties everything between
    so leftover P-XXX / "/" / SCA#- runs disappear.
    """
    def repl(m: re.Match) -> str:
        between = m.group("between")
        # Strip every <w:t>...</w:t> body in the "between" span so the only
        # textual content left in the stamp area is {{REFERENCE}}.
        cleared_between = re.sub(
            r"(<w:t[^>]*>)[^<]*(</w:t>)",
            r"\1\2",
            between,
        )
        return (
            f"{m.group(1)}{{{{REFERENCE}}}}{m.group(2)}"
            f"{cleared_between}"
            f"{m.group(4)}{m.group(5)}"
        )

    return _FOOTER_STAMP_RE.sub(repl, xml_text)


def patch_footer_stamps(docx_path: Path) -> None:
    """Footer XMLs: replace BGCC P-XXX / SCA#-ZZZ stamps with {{REFERENCE}}.

    Done by rewriting the .docx zip in place because python-docx doesn't
    iterate textbox content inside footer xml files (the BGCC stamp sits
    inside a ``<wps:txbx>`` textbox + a ``<v:textbox>`` legacy fallback,
    neither of which is exposed through ``section.footer.paragraphs``).
    """
    tmp_path = docx_path.with_suffix(".docx.rev02_tmp")
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(
        tmp_path, "w", zipfile.ZIP_DEFLATED
    ) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if (
                item.filename.startswith("word/footer")
                and item.filename.endswith(".xml")
            ):
                text = data.decode("utf-8")
                new_text = _rewrite_stamp_in_xml(text)
                data = new_text.encode("utf-8")
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    if not BACKUP_DOCX.exists():
        shutil.copy2(MASTER_DOCX, BACKUP_DOCX)
        print(f"Backup written to {BACKUP_DOCX}")

    doc = Document(str(MASTER_DOCX))
    patch_appendix(doc)
    patch_cover_page_layout(doc)
    patch_remove_form_title_page(doc)
    # Item 11 (3.4(e) page break) intentionally skipped — client reverted
    # the request. The helper `patch_clause_3_4_e_pagebreak` is still
    # defined above so it can be re-enabled by adding a single call here.
    patch_body_stamps(doc)
    doc.save(str(MASTER_DOCX))
    print(
        "python-docx edits applied (Appendix + cover layout "
        "+ section-title page drop + body BGCC stamp → {{REFERENCE}})."
    )

    patch_footer_stamps(MASTER_DOCX)
    print("Footer XMLs patched (BGCC stamp → {{REFERENCE}}).")

    # Confirm tokens are now present.
    doc = Document(str(MASTER_DOCX))
    found = set()
    for p in doc.paragraphs:
        found.update(re.findall(r"\{\{([A-Z][A-Z0-9_]*)\}\}", p.text))
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    found.update(re.findall(r"\{\{([A-Z][A-Z0-9_]*)\}\}", p.text))
    required = {
        "A05", "A06", "A10", "A10_AMOUNT",
        "A16", "A17", "A18",
        "A21", "A21_AMOUNT",
        "REFERENCE",
    }
    missing = required - found
    if missing:
        raise SystemExit(
            f"Patch applied but expected tokens still missing: {sorted(missing)}"
        )
    print(f"Verified tokens present: {sorted(required)}")


if __name__ == "__main__":
    main()
