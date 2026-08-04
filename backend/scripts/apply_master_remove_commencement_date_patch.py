"""Remove the Commencement Date appendix row and its clause-4.3 reference.

Phase 2 req 3 (BGCC "Remove Commencement Date"), scoped minimal per decision:
- Delete Appendix Table 3 row 15 ("Commencement Date" | 4.1 | {{A15}}). Unlike
  the A05 precedent, this row's label is unique (no adjacent duplicate), so it
  can be deleted outright without the row-corruption workaround.
- Rewrite the one clause-4.3 sentence (paragraph ~389) that anchors the
  completion deadline to the Commencement Date, since Time for Completion
  (C08/A17) is already a day-count field, not a date:
    Before: "...within {{C08}} from the Commencement Date or by {{A17}}."
    After:  "...within {{C08}} or by {{A17}}."

NOT touched (out of scope per decision 6): clause 4.1's own definition of
Commencement Date (paras 361-363, 369 — triggered by written instruction, not
by {{A15}}) and the insurance-submission clause (~para 640) — neither reads
the {{A15}} appendix value, so they still function correctly without it.

Idempotent: if {{A15}} is already absent from Table 3, exits without changes.

Run::

    python backend/scripts/apply_master_remove_commencement_date_patch.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"
BACKUP = MASTER_DOCX.with_suffix(".docx.pre-commencement-date.bak")

OLD_SENTENCE_FRAGMENT = "{{C08}} from the Commencement Date or by {{A17}}"
NEW_SENTENCE_FRAGMENT = "{{C08}} or by {{A17}}"


def _token_in_table(table, token: str) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if token in cell.text:
                return True
    return False


def _remove_commencement_date_row(doc) -> bool:
    """Delete the Commencement Date row from Table 3. Returns True if a
    change was made, False if {{A15}} was already absent."""
    appendix_table = doc.tables[3]
    if not _token_in_table(appendix_table, "{{A15}}"):
        return False  # already removed

    target_row = None
    for row in appendix_table.rows:
        if "{{A15}}" in row.cells[-1].text:
            target_row = row
            break

    if target_row is None:
        raise RuntimeError("Could not find {{A15}} row in Table 3")

    target_row._tr.getparent().remove(target_row._tr)
    return True


def _rewrite_clause_4_3_sentence(doc) -> bool:
    """Drop the 'from the Commencement Date' anchor from the completion
    deadline sentence. Returns True if a change was made."""
    changed = False
    for p in doc.paragraphs:
        if OLD_SENTENCE_FRAGMENT in p.text and p.runs:
            # The whole sentence lives in run 0 in the current master;
            # substitute there so we don't disturb paragraph formatting.
            for run in p.runs:
                if OLD_SENTENCE_FRAGMENT in run.text:
                    run.text = run.text.replace(OLD_SENTENCE_FRAGMENT, NEW_SENTENCE_FRAGMENT)
                    changed = True
    return changed


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    row_changed = _remove_commencement_date_row(doc)
    sentence_changed = _rewrite_clause_4_3_sentence(doc)

    if not row_changed and not sentence_changed:
        print("Commencement Date already removed — no change.")
        return

    if not BACKUP.exists():
        shutil.copy2(MASTER_DOCX, BACKUP)

    doc.save(str(MASTER_DOCX))

    doc2 = Document(str(MASTER_DOCX))
    if _token_in_table(doc2.tables[3], "{{A15}}"):
        raise SystemExit("Patch ran but {{A15}} is still present in Table 3.")
    if any(OLD_SENTENCE_FRAGMENT in p.text for p in doc2.paragraphs):
        raise SystemExit("Patch ran but the old Commencement Date sentence is still present.")
    print("Commencement Date appendix row removed; clause 4.3 sentence updated.")


if __name__ == "__main__":
    main()
