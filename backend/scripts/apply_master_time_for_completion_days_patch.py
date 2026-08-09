"""Append " days" to both Time for Completion appendix lines (A16 and A17).

Before:
  Time for Completion of the Project: {{A16}}
  Time for Completion of the Subcontract Works: {{A17}}

After:
  Time for Completion of the Project: {{A16}} days
  Time for Completion of the Subcontract Works: {{A17}} days

Both sentences live in ONE run (a single <w:r> with two <w:t> children
separated by a <w:br/>) inside Appendix Table 3's "Time for Completion"
row — a straightforward text-node edit, no run/paragraph restructuring
needed.

Idempotent: if the row already ends both lines with " days", exits
without changes.

Run::

    python backend/scripts/apply_master_time_for_completion_days_patch.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"
BACKUP = MASTER_DOCX.with_suffix(".docx.pre-tfc-days.bak")


def _find_time_for_completion_row(doc):
    t3 = doc.tables[3]
    for row in t3.rows:
        label = row.cells[0].text
        if "Time for Completion" in label and "Sections" not in label:
            return row
    raise RuntimeError("Could not find the Time for Completion row in Table 3")


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    row = _find_time_for_completion_row(doc)
    value_cell = row.cells[-1]
    para = value_cell.paragraphs[0]
    run0 = para.runs[0]
    t_elements = run0._r.findall(qn("w:t"))

    if len(t_elements) != 2:
        raise SystemExit(f"Expected 2 <w:t> elements in the run, found {len(t_elements)}")

    if t_elements[0].text.endswith(" days") and t_elements[1].text.endswith(" days"):
        print("Time for Completion lines already end with 'days' — no change.")
        return
    if "{{A16}}" not in t_elements[0].text or "{{A17}}" not in t_elements[1].text:
        raise SystemExit(
            "Unexpected text in the Time for Completion run — "
            f"got {t_elements[0].text!r} / {t_elements[1].text!r}"
        )

    t_elements[0].text = t_elements[0].text + " days"
    t_elements[1].text = t_elements[1].text + " days"

    if not BACKUP.exists():
        shutil.copy2(MASTER_DOCX, BACKUP)

    doc.save(str(MASTER_DOCX))

    doc2 = Document(str(MASTER_DOCX))
    row2 = _find_time_for_completion_row(doc2)
    text2 = row2.cells[-1].text
    if "{{A16}} days" not in text2 or "{{A17}} days" not in text2:
        raise SystemExit(f"Patch ran but expected text not found. Got: {text2!r}")
    print("Time for Completion lines now end with 'days'.")


if __name__ == "__main__":
    main()
