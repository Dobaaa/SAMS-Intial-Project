"""Two patches for the master docx.

1. C02 inline in sentence
   Clause 3.3 previously had a standalone highlighted paragraph '{{C02}}' (P257)
   sitting above the sentence 'The Quantities mentioned in this Subcontract
   Agreement is ………………' (P259). This made the rendered PDF show:

       Lump Sum
       The Quantities mentioned in this Subcontract Agreement is ………………

   Fix: replace the dots run in the sentence with '{{C02}}' and delete the
   standalone paragraph so the sentence renders as:

       The Quantities mentioned in this Subcontract Agreement is Lump Sum

2. Remove all yellow highlights
   Every <w:highlight> element in every run-property block (body, tables,
   headers/footers) is removed so no yellow highlighting appears in the PDF.

Both patches are idempotent.

Run::

    python backend/scripts/apply_master_c02_inline_and_dehighlight_patch.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"

# The ellipsis text that marks the blank to be filled by the contract type
DOTS = "………………"


def _remove_all_highlights(doc) -> int:
    """Strip every <w:highlight> element from every run in the document.

    Returns the count of highlight elements removed.
    """
    removed = 0

    def _strip_para(para):
        nonlocal removed
        for run in para.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is None:
                continue
            for hl in rpr.findall(qn("w:highlight")):
                rpr.remove(hl)
                removed += 1

    # Body paragraphs
    for para in doc.paragraphs:
        _strip_para(para)

    # Table cell paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _strip_para(para)

    # Headers / footers
    for section in doc.sections:
        for attr in (
            "header", "first_page_header", "even_page_header",
            "footer", "first_page_footer", "even_page_footer",
        ):
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            for para in getattr(hf, "paragraphs", []):
                _strip_para(para)

    return removed


def _fix_c02_sentence(doc) -> tuple[bool, bool]:
    """
    1. In the body paragraph 'The Quantities mentioned...', replace the DOTS
       run text with '{{C02}}'.
    2. Delete any standalone body paragraph whose full text is exactly '{{C02}}'
       and that still carries a yellow highlight (the now-redundant standalone).

    Returns (sentence_fixed, standalone_deleted).
    """
    sentence_fixed = False
    standalone_deleted = False

    paras = doc.paragraphs

    # --- Fix the sentence ---
    for para in paras:
        if DOTS in para.text and "Quantities mentioned" in para.text:
            for run in para.runs:
                if DOTS in run.text:
                    run.text = run.text.replace(DOTS, "{{C02}}")
                    sentence_fixed = True
            break  # only one such paragraph

    # --- Delete the standalone {{C02}} paragraph (highlighted) ---
    # Match by: text == "{{C02}}" AND has at least one yellow highlight run.
    # Using the XML parent removal approach (python-docx doesn't expose
    # paragraph deletion directly).
    for para in list(paras):
        if para.text.strip() != "{{C02}}":
            continue
        has_yellow = any(
            run._r.find(qn("w:rPr")) is not None
            and run._r.find(qn("w:rPr")).find(qn("w:highlight")) is not None
            for run in para.runs
        )
        if has_yellow:
            p_el = para._element
            p_el.getparent().remove(p_el)
            standalone_deleted = True
            break

    return sentence_fixed, standalone_deleted


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))

    sentence_fixed, standalone_deleted = _fix_c02_sentence(doc)
    highlights_removed = _remove_all_highlights(doc)

    any_change = sentence_fixed or standalone_deleted or highlights_removed > 0
    if not any_change:
        print("All patches already applied — no change.")
        return

    if sentence_fixed:
        print("C02 sentence: dots replaced with {{C02}} token.")
    else:
        print("C02 sentence: already contains {{C02}} (no change).")

    if standalone_deleted:
        print("Standalone {{C02}} paragraph (highlighted) deleted.")
    else:
        print("Standalone {{C02}} paragraph: not found or already removed.")

    print(f"Highlights removed: {highlights_removed}")

    doc.save(str(MASTER_DOCX))

    # Verify: no highlights remain, sentence has the token
    doc2 = Document(str(MASTER_DOCX))
    hl_count = sum(
        1
        for para in list(doc2.paragraphs) + [
            p for tbl in doc2.tables for row in tbl.rows for cell in row.cells for p in cell.paragraphs
        ]
        for run in para.runs
        if run._r.find(qn("w:rPr")) is not None
        and run._r.find(qn("w:rPr")).find(qn("w:highlight")) is not None
    )
    if hl_count:
        raise SystemExit(f"Verification failed: {hl_count} highlights still present.")

    sentence_ok = any(
        "{{C02}}" in para.text and "Quantities mentioned" in para.text
        for para in doc2.paragraphs
    )
    if not sentence_ok:
        raise SystemExit("Verification failed: sentence does not contain {{C02}}.")

    print("Verification passed.")


if __name__ == "__main__":
    main()
