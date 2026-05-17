"""One-shot: turn the user's fully-populated sample .docx into a SAMS master
template by replacing each demo agreement value with a {{FIELD_ID}} token.

Usage:
    python -m scripts.tokenize_master_docx /home/seif/Downloads/sams_tahoma.docx

Writes to backend/masters/sca_master_v1.docx (overwriting). The resulting
docx is what services.docx_pdf_service loads at render time.

The demo data this script targets matches what we seeded for the
Marina Tower Phase II / Elite Steel Industries demo on alpha. Anyone
re-running this script must first re-create the docx by rendering an
agreement with those exact values, OR update VALUE_TO_TOKEN to whatever
the new sample docx contains.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph

BACKEND_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = BACKEND_DIR / "masters" / "sca_master_v1.docx"


# Demo value -> SAMS token. Order matters: longer strings replaced first so
# they don't get clobbered by shorter substrings of themselves.
VALUE_TO_TOKEN: list[tuple[str, str]] = [
    # Project / subcontractor identification
    ("Marina Tower Phase II — Structural Steel Package", "{{F06}}"),
    ("Structural Steel Fabrication, Surface Treatment, and Erection", "{{F09}}"),
    ("Dubai Holding Investments LLC", "{{F05}}"),
    ("Atkins Middle East Consulting", "{{A03}}"),
    ("Elite Steel Industries LLC", "{{F02}}"),
    ("Plot 27, Dubai Marina, UAE", "{{F07}}"),
    ("TL-DXB-2024-118503", "{{F04}}"),
    # Payment + retention
    ("12 months after Taking-Over", "{{C07}}"),
    ("Upon Taking-Over Certificate", "{{C06}}"),
    ("Against bank guarantee", "{{A11}}"),
    ("Dubai Courts, UAE", "{{C13}}"),
    # Dates and numbers
    ("22nd May 2026", "{{F01}}"),  # already long-formatted in the sample
    ("60 days PDC", "{{C05}}"),
    ("8,500,000.00", "{{F08}}"),
    ("850,000.00", "{{C03}}"),
    ("12,000.00", "{{C11}}"),
    ("Lump Sum", "{{C02}}"),
    ("18 months", "{{C08}}"),
    ("15 months", "{{A17}}"),
    ("2026-06-01", "{{A15}}"),
    ("78421", "{{F03}}"),
]

# Some appendix cells contain a bare number ("12" for DLP, "14" for Insurance
# days). String-replacing "12" or "14" globally would catch unrelated
# substrings (page numbers, "12 months" prefix, etc.), so we patch these
# row-by-row keyed on the row label in column 0.
APPENDIX_ROW_TOKEN: dict[str, str] = {
    "Defects Liability Period": "{{A19}}",
    "Time to submit the Copies of the required Insurance Policies": "{{A22}}",
}


# Context-anchored token rewrites for the legacy "(……Insert…..)" prose
# markers that appear inside the Conditions text but were not normalised
# to {{C##}} tokens. Each tuple is
# (substring_in_paragraph_text, occurrence_index, token_to_insert).
# occurrence_index selects WHICH of the parenthetical Insert markers in
# that paragraph gets replaced (0-based) so a paragraph with two
# placeholders can have each one mapped to a different field.
PARENTHETICAL_INSERT = "(……Insert…..)"

CONTEXT_INSERT_TOKEN: list[tuple[str, int, str]] = [
    # 3.4.1 Advance Payment release condition
    ("Such advance payment shall be released by the Main Contractor", 0, "{{C04}}"),
    # 3.4.6 Interim Payment days
    ("amounts properly due under this Subcontract within", 0, "{{C05}}"),
    # 3.4.7 1st Half Retention release days
    ("First instalment (5%) to be released within", 0, "{{C06}}"),
    # 3.4.7 2nd Half Retention release days
    ("The Second instalment (5%) to be released", 0, "{{C07}}"),
    # 4.3 Time for Completion — Project (months) AND completion date
    ("complete the entire Subcontract Works", 0, "{{C08}}"),
    ("complete the entire Subcontract Works", 1, "{{A17}}"),
    # 5 Defects Liability Period (months)
    ("Months the Defects Liability Period (DLP)", 0, "{{C10}}"),
    # 6.2 Rate of Liquidated Damages (AED/day)
    ("UAE per calendar day in part or full day", 0, "{{C11}}"),
    # 10.1 Insurance submission deadline (days from commencement)
    ("from the Commencement Date submit to the Main Contractor a copy of the required", 0, "{{C12}}"),
    # 13.3 Dispute Resolution Jurisdiction
    ("United Arab Emirates, which shall have full", 0, "{{C13}}"),
]


def _iter_paragraphs(doc: DocxDocument):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _apply_contextual_inserts(
    doc: DocxDocument,
    rules: list[tuple[str, int, str]],
) -> int:
    """Replace specific occurrences of "(……Insert…..)" by paragraph context.

    For each (context_substring, occurrence_index, token) tuple, find a
    paragraph that contains both the context substring AND the literal
    "(……Insert…..)" marker, then replace the Nth occurrence (0-based)
    of that marker with the token. Returns the number of replacements
    actually applied. Missing matches are reported so we can iterate."""
    # Group rules by context_substring so we can do per-paragraph
    # multi-replacement in one pass.
    by_ctx: dict[str, list[tuple[int, str]]] = {}
    for ctx, idx, token in rules:
        by_ctx.setdefault(ctx, []).append((idx, token))

    applied = 0
    for ctx, ordered_rules in by_ctx.items():
        target = None
        for p in _all_paragraphs(doc):
            text = p.text or ""
            if ctx in text and PARENTHETICAL_INSERT in text:
                target = p
                break
        if target is None:
            print(f"    ! context not found: {ctx[:60]!r}")
            continue
        full = "".join(r.text or "" for r in target.runs)
        # Sort by occurrence index so replacements happen left-to-right.
        ordered_rules.sort()
        new_text = full
        cursor = 0
        for idx, token in ordered_rules:
            # Find the (idx+1)-th occurrence starting from `cursor`
            pos = -1
            occurrences_seen = 0
            search_from = cursor
            while True:
                hit = new_text.find(PARENTHETICAL_INSERT, search_from)
                if hit == -1:
                    break
                if occurrences_seen == 0 and pos == -1:
                    pos = hit
                occurrences_seen += 1
                if occurrences_seen > idx:
                    pos = hit
                    break
                search_from = hit + len(PARENTHETICAL_INSERT)
            if pos == -1:
                print(f"    ! marker[{idx}] missing in: {ctx[:60]!r}")
                continue
            new_text = new_text[:pos] + token + new_text[pos + len(PARENTHETICAL_INSERT):]
            cursor = pos + len(token)
            applied += 1
        # Collapse into the first run
        if target.runs:
            target.runs[0].text = new_text
            for r in target.runs[1:]:
                r.text = ""
    return applied


def _all_paragraphs(doc):
    """Body + table-cell paragraphs in document order."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _split_cover_paragraph(doc: DocxDocument) -> None:
    """Find the concatenated cover paragraph and split it into 4 rows."""
    from copy import deepcopy
    from lxml import etree

    # The marker text only appears in the one concatenated paragraph.
    marker = "Scope Title"
    target_para: Paragraph | None = None
    for para in doc.paragraphs:
        if marker in para.text and "Subcontractor Name" in para.text:
            target_para = para
            break
    if target_para is None:
        print("  ! cover concatenation not detected — no split performed")
        return

    # The four lines we want, each preserving the tab between label and value.
    # Colons restored to match the BGCC source PDF (the docx revision the user
    # provided dropped them when going through whatever round-trip happened).
    new_lines = [
        "The Project\t: {{F06}}",
        "Subcontractor Name\t: {{F02}}",
        "Scope Title\t: {{F09}}",
        "Subcontract Agreement Documents:",
    ]

    # Copy the pPr (paragraph properties) and rPr from the original so the
    # split paragraphs keep tab stops, indent, color, font.
    orig_el = target_para._element
    orig_pPr = orig_el.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
    )
    # Reuse the original paragraph's first run's properties for the new runs.
    src_run = target_para.runs[0]._element
    src_rPr = src_run.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
    )

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def make_paragraph(text: str):
        p = etree.SubElement(orig_el.getparent(), f"{{{W}}}p")
        if orig_pPr is not None:
            p.append(deepcopy(orig_pPr))
        # Split on tab so we emit <w:tab/> elements, not literal "\t"
        parts = text.split("\t")
        for i, part in enumerate(parts):
            r = etree.SubElement(p, f"{{{W}}}r")
            if src_rPr is not None:
                r.append(deepcopy(src_rPr))
            if part:
                t = etree.SubElement(r, f"{{{W}}}t")
                t.text = part
                t.set(
                    "{http://www.w3.org/XML/1998/namespace}space",
                    "preserve",
                )
            if i < len(parts) - 1:
                etree.SubElement(r, f"{{{W}}}tab")
        return p

    # Insert the new paragraphs right after the original, then remove it.
    parent = orig_el.getparent()
    idx = list(parent).index(orig_el)
    for offset, line in enumerate(new_lines):
        new_p = make_paragraph(line)
        # SubElement appended at end; move it to the right slot.
        parent.remove(new_p)
        parent.insert(idx + 1 + offset, new_p)
    parent.remove(orig_el)
    print(f"  ✓ split cover paragraph into {len(new_lines)} rows")


def _replace_across_runs(para: Paragraph, old: str, new: str) -> bool:
    """Replace `old` with `new` in a paragraph even when run boundaries split
    the match. Returns True if a substitution happened.

    Collapses run text into the first run; loses per-run inline formatting
    within the matched span (acceptable for plain-text demo values being
    replaced by plain-text tokens)."""
    if not para.runs:
        return False
    full = "".join(r.text or "" for r in para.runs)
    if old not in full:
        return False
    para.runs[0].text = full.replace(old, new)
    for run in para.runs[1:]:
        run.text = ""
    return True


def main(input_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    doc = Document(str(input_path))

    print(f"Source: {input_path}  ({sum(1 for _ in _iter_paragraphs(doc))} paragraphs)")

    # 1. General value->token replacements.
    replaced: dict[str, int] = {}
    for old, token in VALUE_TO_TOKEN:
        count = 0
        for para in _iter_paragraphs(doc):
            if _replace_across_runs(para, old, token):
                count += 1
        replaced[token] = count

    # 2. Row-label-keyed replacements for bare-number cells.
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 3:
                continue
            label = row.cells[0].text.strip()
            token = APPENDIX_ROW_TOKEN.get(label)
            if not token:
                continue
            cell = row.cells[2]
            for para in cell.paragraphs:
                if para.text.strip() and para.text.strip().isdigit():
                    if para.runs:
                        para.runs[0].text = token
                        for r in para.runs[1:]:
                            r.text = ""
                    replaced[token] = replaced.get(token, 0) + 1
                    break  # only the first numeric paragraph

    # The user-edited 42-page docx collapsed the three cover label rows
    # (Project / Subcontractor Name / Scope Title) plus the "Subcontract
    # Agreement Documents:" header into a single paragraph separated by
    # tabs and spaces. LibreOffice renders that single paragraph as a
    # wrapped block, not three left-aligned rows like the source PDF.
    # Split it back into four paragraphs with the same paragraph
    # formatting so the cover renders correctly.
    _split_cover_paragraph(doc)

    # 3. Replace "(……Insert…..)" markers using surrounding-text anchors.
    # Each paragraph keeps its formatting; we just swap the literal
    # marker for the right {{C##}} token. Indices are needed because
    # one paragraph (Time for Completion) carries two placeholders.
    insert_replaced = _apply_contextual_inserts(doc, CONTEXT_INSERT_TOKEN)
    print(f"  ✓ contextual Insert markers tokenized: {insert_replaced}")

    # 4. Any remaining bare "(……Insert…..)" cells (e.g. the Milestones
    # table's free-entry cells) are blanked out so the rendered PDF
    # doesn't show the literal marker. Admin enters per-agreement
    # milestone values via the field editor.
    cleared = 0
    for p in _all_paragraphs(doc):
        if p.text.strip() == PARENTHETICAL_INSERT and p.runs:
            p.runs[0].text = ""
            for r in p.runs[1:]:
                r.text = ""
            cleared += 1
    if cleared:
        print(f"  ✓ bare Insert cells blanked: {cleared}")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))

    print(f"\nReplacement summary ({len(replaced)} tokens):")
    for token, count in sorted(replaced.items()):
        mark = "✓" if count > 0 else "✗"
        print(f"  {mark} {token:10s}  x{count}")
    print(f"\nWrote tokenized master -> {OUTPUT_PATH}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "input",
        nargs="?",
        default="/home/seif/Downloads/sams_tahoma.docx",
        help="Source docx (populated with demo values) to tokenize",
    )
    args = parser.parse_args()
    main(Path(args.input))
