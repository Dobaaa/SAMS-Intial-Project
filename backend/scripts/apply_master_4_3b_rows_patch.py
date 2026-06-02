"""Add A24/A25/A26 rows to the 4.3(b) Time for Completion table in the SCA master.

BGCC requested three new rows immediately after the existing A18 (Milestones)
row in the page-22 Appendix table under clause 4.3(b):

  Start of Material Submission         {{A24}}
  Complete all Material Submission     {{A25}}
  Start of Submission of Shop Drawings {{A26}}

This script locates the row whose value cell contains ``{{A18}}``, clones it
three times, and inserts the new rows directly after it in the same table.

Idempotent: if ``{{A24}}`` is already present anywhere in the document the
script exits without modifying the file.

Run::

    python backend/scripts/apply_master_4_3b_rows_patch.py
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"

_NEW_ROWS = [
    ("Start of Material Submission", "{{A24}}"),
    ("Complete all Material Submission", "{{A25}}"),
    ("Start of Submission of Shop Drawings", "{{A26}}"),
]


def _set_cell_text(cell, text: str) -> None:
    """Replace cell content with `text`, preserving first-paragraph run formatting."""
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _doc_has_token(doc, token: str) -> bool:
    for p in doc.paragraphs:
        if token in p.text:
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if token in cell.text:
                    return True
    return False


def patch_4_3b_rows(doc) -> bool:
    """Insert A24/A25/A26 rows after the A18 Milestones row.

    Returns True if rows were added, False if already present.
    """
    if _doc_has_token(doc, "{{A24}}"):
        return False

    # Find the table and row whose last cell contains {{A18}}.
    target_table = None
    a18_tr = None
    for table in doc.tables:
        for row in table.rows:
            if "{{A18}}" in row.cells[-1].text or "{{A18}}" in row.cells[2].text if len(row.cells) > 2 else False:
                target_table = table
                a18_tr = row._tr
                break
        if target_table is not None:
            break

    if target_table is None:
        # Fallback: search all cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{A18}}" in cell.text:
                        target_table = table
                        a18_tr = row._tr
                        break
                if a18_tr is not None:
                    break
            if a18_tr is not None:
                break

    if a18_tr is None:
        raise RuntimeError(
            "Could not find the {{A18}} row in the master docx. "
            "Run apply_master_rev02_patches.py first."
        )

    # Insert new rows in reverse so each addnext places them right after A18
    # in the correct final order: A24, A25, A26.
    for label, token in reversed(_NEW_ROWS):
        new_tr = deepcopy(a18_tr)
        a18_tr.addnext(new_tr)

        # The newly inserted row is now the sibling immediately after a18_tr.
        # Retrieve it via the table's row list — it is the row whose _tr is new_tr.
        new_row = None
        for row in target_table.rows:
            if row._tr is new_tr:
                new_row = row
                break

        if new_row is None:
            raise RuntimeError("Inserted row not found — unexpected table structure.")

        _set_cell_text(new_row.cells[0], label)
        # cells[1] is the clause-ref column; it inherits "4.3(b)" from the A18
        # clone which is correct — leave it unchanged.
        value_col = 2 if len(new_row.cells) > 2 else len(new_row.cells) - 1
        _set_cell_text(new_row.cells[value_col], token)

    return True


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    bak = MASTER_DOCX.with_suffix(".docx.pre-4_3b.bak")
    if not bak.exists():
        import shutil
        shutil.copy2(MASTER_DOCX, bak)
        print(f"Backup written to {bak.name}")

    doc = Document(str(MASTER_DOCX))
    added = patch_4_3b_rows(doc)
    if not added:
        print("{{A24}} already present — no change.")
        return

    doc.save(str(MASTER_DOCX))

    # Verify all three tokens landed.
    doc = Document(str(MASTER_DOCX))
    for token in ("{{A24}}", "{{A25}}", "{{A26}}"):
        if not _doc_has_token(doc, token):
            raise SystemExit(f"Patch ran but {token} is still missing — check the master.")
    print("Added rows A24/A25/A26 to the 4.3(b) table in the master docx.")


if __name__ == "__main__":
    main()
