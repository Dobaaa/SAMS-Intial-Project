"""Wire {{A24}}/{{A25}}/{{A26}} tokens into the page-22 milestone table (Table 5).

The milestone table on page 22 already has three labelled rows with empty
"Time for Completion" cells:

  Start of Material Submission          → {{A24}}
  Complete all Material Submission      → {{A25}}
  Start of Submission of Shop Drawings  → {{A26}}

This script also removes the three A24/A25/A26 rows that were previously
inserted into Table 3 (the appendix summary table) by mistake, so those
tokens only appear in the correct milestone table.

Idempotent: if {{A24}} is already present in Table 5 (the milestone table)
the script exits without modifying the file.

Run::

    python backend/scripts/apply_master_milestone_table_patch.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"

# Maps a substring of the milestone row label (lower-case) → token to place
# in cells[1] (the "Time for Completion" column).
_MILESTONE_MAP = [
    ("start of material submission", "{{A24}}"),
    ("complete all material", "{{A25}}"),
    ("shop", "{{A26}}"),
]

# Tokens that should NOT appear in Table 3 (appendix summary).
_APPENDIX_TOKENS_TO_REMOVE = {"{{A24}}", "{{A25}}", "{{A26}}"}


def _set_cell_text(cell, text: str) -> None:
    """Replace cell content preserving the first paragraph's run formatting."""
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _cell_contains_any(cell, tokens: set[str]) -> bool:
    return any(t in cell.text for t in tokens)


def _token_in_table(table, token: str) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if token in cell.text:
                return True
    return False


def remove_wrong_appendix_rows(doc) -> int:
    """Remove any A24/A25/A26 rows from Table 3 (appendix summary).

    Returns the number of rows removed.
    """
    appendix_table = doc.tables[3]
    removed = 0
    for row in list(appendix_table.rows):
        if _cell_contains_any(row.cells[-1], _APPENDIX_TOKENS_TO_REMOVE):
            row._tr.getparent().remove(row._tr)
            removed += 1
    return removed


def patch_milestone_table(doc) -> bool:
    """Add tokens to the Time-for-Completion cells in Table 5.

    Returns True if changes were made, False if already patched.
    """
    milestone_table = doc.tables[5]

    # Idempotency: if {{A24}} is already in Table 5 we're done.
    if _token_in_table(milestone_table, "{{A24}}"):
        return False

    for row in milestone_table.rows:
        label = row.cells[0].text.strip().lower()
        for keyword, token in _MILESTONE_MAP:
            if keyword in label:
                _set_cell_text(row.cells[1], token)
                break

    return True


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))

    # Step 1: remove the wrong appendix-summary rows (idempotent — no-op if
    # they're already gone).
    removed = remove_wrong_appendix_rows(doc)
    if removed:
        print(f"Removed {removed} A24/A25/A26 row(s) from the appendix summary table.")
    else:
        print("Appendix summary table already clean — no rows to remove.")

    # Step 2: add tokens to the milestone table.
    patched = patch_milestone_table(doc)
    if not patched:
        print("Milestone table already has {{A24}} — token step skipped.")
    else:
        doc.save(str(MASTER_DOCX))
        # Verify
        doc2 = Document(str(MASTER_DOCX))
        for token in ("{{A24}}", "{{A25}}", "{{A26}}"):
            if not _token_in_table(doc2.tables[5], token):
                raise SystemExit(f"Patch ran but {token} is missing from Table 5.")
        print("Wired {{A24}}/{{A25}}/{{A26}} into the page-22 milestone table.")
        return

    # If we only removed rows (step 1) and didn't need step 2, still save.
    if removed:
        doc.save(str(MASTER_DOCX))
        print("Saved cleaned master docx.")


if __name__ == "__main__":
    main()
