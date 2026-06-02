"""Collapse the Communications Address section to Subcontractor-only.

Before (two rows):
  Communications Address for Serving of the Notices | 1.6 | The Main Contractor Address\n{{A05}}
  Communications Address for Serving of the Notices | 1.6 | The Subcontractor Address\n{{A06}}

After (one row):
  Communications Address for Serving of the Notices | 1.6 | The Subcontractor Address\n{{A06}}

Strategy: overwrite the A05 row's value cell with the Subcontractor content,
then delete the A06 row. This preserves the label cell intact (removing the
first of two identically-labelled rows corrupts the second row's label in
python-docx due to shared XML nodes).

Idempotent: if {{A05}} is already absent from Table 3 the script exits without
modifying the file.

Run::

    python backend/scripts/apply_master_remove_a05_row_patch.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"


def _set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _token_in_table(table, token: str) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if token in cell.text:
                return True
    return False


def collapse_comms_address(doc) -> bool:
    """Merge the two Communications Address rows into one (Subcontractor only).

    Returns True if changes were made, False if already collapsed.
    """
    appendix_table = doc.tables[3]
    if not _token_in_table(appendix_table, "{{A05}}"):
        return False  # already done

    a05_row = None
    a06_row = None
    for row in appendix_table.rows:
        value_text = row.cells[-1].text
        if "{{A05}}" in value_text:
            a05_row = row
        elif "{{A06}}" in value_text and a05_row is not None:
            a06_row = row
            break

    if a05_row is None:
        raise RuntimeError("Could not find {{A05}} row in Table 3")
    if a06_row is None:
        raise RuntimeError("Found {{A05}} but could not find following {{A06}} row")

    # Repurpose the A05 row: overwrite its value cell with the Subcontractor
    # content so the label cell stays intact.
    value_col = 2 if len(a05_row.cells) > 2 else len(a05_row.cells) - 1
    _set_cell_text(a05_row.cells[value_col], "The Subcontractor Address\n{{A06}}")

    # Now delete the original A06 row (it's no longer needed).
    a06_row._tr.getparent().remove(a06_row._tr)

    return True


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    changed = collapse_comms_address(doc)
    if not changed:
        print("{{A05}} already absent — no change.")
        return

    doc.save(str(MASTER_DOCX))

    doc2 = Document(str(MASTER_DOCX))
    if _token_in_table(doc2.tables[3], "{{A05}}"):
        raise SystemExit("Patch ran but {{A05}} is still present.")
    if not _token_in_table(doc2.tables[3], "{{A06}}"):
        raise SystemExit("Patch ran but {{A06}} is missing — unexpected.")
    print("Communications Address section collapsed to Subcontractor-only ({{A06}}).")


if __name__ == "__main__":
    main()
