"""Client feedback round (2026-08-24): three changes to the master docx.

1. Commencement Date (Clause 4.1 / A15) removed entirely — not just hidden.
   Clause 4.1's own definition (heading + a/b/c) is deleted from the
   Conditions document, its Appendix Table 3 row is deleted, and every
   other clause that anchored a deadline to "the Commencement Date" is
   reworded to anchor to "the date of this Subcontract Agreement" instead:
   4.2(a) (work programme submission), 4.3(a) (Time for Completion — now
   ONLY the Subcontract Works duration via A17, the Project-level A16/C08
   line is dropped per item 2), and the insurance-policy submission clause.
   Because clauses renumber downward (4.2->4.1, 4.3->4.2, ...) the literal
   clause-ref text in Table 3 for the Time for Completion / Milestones rows
   is updated 4.3(a)->4.2(a) / 4.3(b)->4.2(b) to match (auto-numbered
   headings in the body re-number themselves via Word's numPr, no edit
   needed there).
2. Time for Completion of the Project (A16, auto-sourced from C08) removed
   from the Appendix — only "Time for Completion of the Subcontract Works"
   (A17) remains in that row. A16/C08's DB visibility is handled by
   migration 024, not this script.
3. 3.4.6(a) / 3.4.7(a)(I)/(II): "60-day Post-Dated Cheque (PDC)" -> "30-day"
   per client-confirmed correction (15-day Engineer's Certificate window
   unchanged). Also removes the stray leftover empty paragraphs in that
   region (left over from Package G's clause rewrite) that were causing
   the irregular page-21 vertical gaps the client flagged.

Also removes the stale "Commencement Date" line from the document's frozen
Table of Contents. That TOC is NOT a live field — it's cached paragraphs
wrapped in a content-control (``w:sdt``) that ``python-docx``'s normal
``doc.paragraphs`` does not walk, so it's addressed separately via a
raw XML search over every ``w:p`` in the body (``doc.element.body.iter``).
Confirmed via a real render that the TOC's clause numbering (4.2->4.1 etc.)
re-numbers itself automatically once the stale line is removed, same as
the body headings.

Idempotent: each of the three fixes above checks its own precondition and
is skipped if already applied.

Run::

    python backend/scripts/apply_master_remove_commencement_time_pdc_patch.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"
BACKUP = MASTER_DOCX.with_suffix(".docx.pre-commencement-pdc.bak")


def _set_cell_text(cell, text: str) -> None:
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _set_paragraph_text(para, text: str) -> None:
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _find_paragraph(doc, token: str):
    for p in doc.paragraphs:
        if token in p.text:
            return p
    raise RuntimeError(f"Could not find a paragraph containing {token!r}")


def _delete_paragraph(para) -> None:
    para._element.getparent().remove(para._element)


def _token_in_table(table, token: str) -> bool:
    return any(token in cell.text for row in table.rows for cell in row.cells)


def _remove_clause_4_1(doc) -> None:
    heading = next(p for p in doc.paragraphs if p.text.strip() == "Commencement Date")

    to_delete = [heading]
    node = heading._p.getnext()
    while node is not None:
        cand = Paragraph(node, heading._parent)
        if cand.text.strip() == "Programme of the Subcontract Works":
            break
        to_delete.append(cand)
        node = node.getnext()

    prev_node = heading._p.getprevious()
    if prev_node is not None:
        prev_para = Paragraph(prev_node, heading._parent)
        if not prev_para.text.strip():
            to_delete.insert(0, prev_para)

    for p in to_delete:
        _delete_paragraph(p)


def _reword_commencement_date_references(doc) -> None:
    p1 = _find_paragraph(doc, "Within seven (7) days from the Commencement Date,")
    _set_paragraph_text(
        p1,
        p1.text.replace(
            "Within seven (7) days from the Commencement Date,",
            "Within seven (7) days from the date of this Subcontract Agreement,",
        ),
    )

    p2 = _find_paragraph(doc, "within {{C08}} from the Commencement Date or by {{A17}}.")
    _set_paragraph_text(
        p2,
        p2.text.replace(
            "within {{C08}} from the Commencement Date or by {{A17}}.",
            "within {{A17}} days from the date of this Subcontract Agreement.",
        ),
    )

    p3 = _find_paragraph(doc, "within {{C12}} from the Commencement Date submit")
    _set_paragraph_text(
        p3,
        p3.text.replace(
            "within {{C12}} from the Commencement Date submit",
            "within {{C12}} from the date of this Subcontract Agreement submit",
        ),
    )


def _update_table3(doc) -> None:
    table3 = doc.tables[3]

    a15_row = next(row for row in table3.rows if "{{A15}}" in row.cells[-1].text)
    a15_row._tr.getparent().remove(a15_row._tr)

    toc_row = next(row for row in table3.rows if "{{A16}}" in row.cells[-1].text)
    _set_cell_text(toc_row.cells[-1], "Time for Completion of the Subcontract Works: {{A17}} days")
    _set_cell_text(toc_row.cells[1], "4.2(a)")

    milestones_row = next(row for row in table3.rows if "{{A18}}" in row.cells[-1].text)
    _set_cell_text(milestones_row.cells[1], "4.2(b)")


def _fix_pdc_wording_and_gaps(doc) -> None:
    p311 = _find_paragraph(doc, "Progress payments shall be released by 60-day")
    _set_paragraph_text(
        p311, p311.text.replace("60-day", "30-day").replace("60 days", "30 days")
    )

    for _ in range(2):
        p = _find_paragraph(doc, "60-day Post-Dated Cheque (PDC)")
        _set_paragraph_text(p, p.text.replace("60-day", "30-day"))

    # Trim the 3 stray blanks before "Retention Money and Final Payment" to 1.
    anchor_c = _find_paragraph(
        doc,
        "shall be entitled to withhold or deduct such amounts from any interim or final",
    )
    blanks = []
    node = anchor_c._p.getnext()
    while node is not None:
        cand = Paragraph(node, anchor_c._parent)
        if cand.text.strip():
            break
        blanks.append(cand)
        node = node.getnext()
    for b in blanks[:-1]:
        _delete_paragraph(b)

    # Trim the 2 stray blanks between retention item II and item (d) to 0.
    anchor_ii = _find_paragraph(doc, "The remaining 50% of the retention shall be released")
    blanks2 = []
    node = anchor_ii._p.getnext()
    while node is not None:
        cand = Paragraph(node, anchor_ii._parent)
        if cand.text.strip():
            break
        blanks2.append(cand)
        node = node.getnext()
    for b in blanks2:
        _delete_paragraph(b)


def _rewire_pdc_day_inputs(doc) -> bool:
    """Re-insert {{C05}}/{{C06}}/{{C07}} as the PDC day count in
    3.4.6(a)/3.4.7(a)(I)/(II) so the admin's Interim/Retention day-count
    inputs actually drive the rendered PDF (2026-08-24 follow-up: the day
    count must stay an admin input, not a hardcoded 30/60 — the earlier fix
    in this same round only swapped one hardcoded number for another).
    The 15-day Engineer's-Certificate-approval window stays fixed text,
    per client confirmation. Returns True if anything changed."""
    changed = False

    p1 = _find_paragraph(doc, "Progress payments shall be released by")
    if "{{C05}}" not in p1.text:
        new_text = p1.text
        for old, new in (
            ("released by 30-day Post-Dated Cheque (PDC)", "released by {{C05}}-day Post-Dated Cheque (PDC)"),
            ("released by 60-day Post-Dated Cheque (PDC)", "released by {{C05}}-day Post-Dated Cheque (PDC)"),
            ("PDC dated 30 days from the invoice date", "PDC dated {{C05}} days from the invoice date"),
            ("PDC dated 60 days from the invoice date", "PDC dated {{C05}} days from the invoice date"),
        ):
            new_text = new_text.replace(old, new)
        _set_paragraph_text(p1, new_text)
        changed = True

    p2 = _find_paragraph(doc, "The first 50% of the retention shall be released")
    if "{{C06}}" not in p2.text:
        new_text = p2.text
        for old in ("30-day", "60-day"):
            new_text = new_text.replace(f"released by {old} Post-Dated Cheque (PDC)", "released by {{C06}}-day Post-Dated Cheque (PDC)")
        _set_paragraph_text(p2, new_text)
        changed = True

    p3 = _find_paragraph(doc, "The remaining 50% of the retention shall be released")
    if "{{C07}}" not in p3.text:
        new_text = p3.text
        for old in ("30-day", "60-day"):
            new_text = new_text.replace(f"released by {old} Post-Dated Cheque (PDC)", "released by {{C07}}-day Post-Dated Cheque (PDC)")
        _set_paragraph_text(p3, new_text)
        changed = True

    return changed


def _all_body_text(doc) -> str:
    """Every ``w:t`` text run anywhere under the body, including ones nested
    inside content controls (``w:sdt``) — e.g. the document's frozen Table
    of Contents — that ``doc.paragraphs`` does not walk."""
    parts = []
    for p in doc.element.body.iter(qn("w:p")):
        texts = p.findall(".//" + qn("w:t"))
        parts.append("".join(t.text or "" for t in texts))
    return "\n".join(parts)


def _find_stale_toc_paragraph(doc):
    for p in doc.element.body.iter(qn("w:p")):
        texts = p.findall(".//" + qn("w:t"))
        text = "".join(t.text or "" for t in texts)
        if "Commencement" in text:
            return p
    return None


def _remove_stale_toc_entry(doc) -> None:
    p = _find_stale_toc_paragraph(doc)
    if p is None:
        raise RuntimeError("Expected a stale TOC 'Commencement Date' paragraph but found none.")
    p.getparent().remove(p)


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    changed = False

    if _token_in_table(doc.tables[3], "{{A15}}"):
        _remove_clause_4_1(doc)
        _reword_commencement_date_references(doc)
        _update_table3(doc)
        _fix_pdc_wording_and_gaps(doc)
        changed = True

    if _find_stale_toc_paragraph(doc) is not None:
        _remove_stale_toc_entry(doc)
        changed = True

    if _rewire_pdc_day_inputs(doc):
        changed = True

    if not changed:
        print("Already patched — no change.")
        return

    if not BACKUP.exists():
        shutil.copy2(MASTER_DOCX, BACKUP)

    doc.save(str(MASTER_DOCX))

    doc2 = Document(str(MASTER_DOCX))
    full_text = _all_body_text(doc2)
    if "Commencement" in full_text:
        raise SystemExit("Patch ran but 'Commencement' text still referenced somewhere (incl. TOC).")
    if "60-day" in full_text or "60 days" in full_text or "30-day" in full_text or "30 days" in full_text:
        raise SystemExit("Patch ran but a hardcoded PDC day count remains (should be a token now).")
    if _token_in_table(doc2.tables[3], "{{A15}}") or _token_in_table(doc2.tables[3], "{{A16}}"):
        raise SystemExit("Patch ran but {{A15}}/{{A16}} still present in Table 3.")
    for tok in ("{{C05}}", "{{C06}}", "{{C07}}"):
        if tok not in full_text:
            raise SystemExit(f"Patch ran but {tok} is missing — PDC day-count rewire failed.")
    print(
        "Commencement Date removed (incl. stale TOC entry), Time for Completion simplified, "
        "PDC day counts rewired to C05/C06/C07 inputs, page-21 gaps fixed."
    )


if __name__ == "__main__":
    main()
