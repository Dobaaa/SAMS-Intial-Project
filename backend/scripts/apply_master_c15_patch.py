"""Append the C15 "Optional Terms" row to the SCA master Appendix.

BGCC asked for a new free-text clause **C15 "Optional Terms"** to render as
the last row of the "APPENDIX TO THE SUBCONTRACT AGREEMENT" continuation
table on page 7 — i.e. directly after the "Dispute Resolution
(Jurisdiction)" row. This script adds that row to
``backend/masters/sca_master_v1.docx`` in one reproducible, idempotent pass:

* Locate the appendix table that ends with the "Dispute Resolution
  (Jurisdiction)" row (by content, so it survives table-index shifts).
* Deep-clone that last ``<w:tr>`` so the new row inherits the table's cell
  borders / Tahoma font / paragraph styling verbatim.
* Rewrite the cloned cells to ``Optional Terms`` | (blank clause column) |
  ``{{C15}}``. The render service substitutes ``{{C15}}`` with the
  agreement's stored value, converting newlines to ``<w:br/>`` so the
  paragraph input renders multi-line.

Idempotent: if ``{{C15}}`` is already present anywhere in the document the
script is a no-op, so re-running after a fresh ``git`` checkout (or on the
VPS) won't duplicate the row.

Run::

    python backend/scripts/apply_master_c15_patch.py
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"

_ROW_LABEL = "Optional Terms"
_ROW_TOKEN = "{{C15}}"


def _set_cell_text(cell, text: str) -> None:
    """Replace cell content with `text`, preserving the first paragraph's
    run-formatting (Tahoma face/size/alignment) and clearing extra runs."""
    paras = cell.paragraphs
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for extra in paras[1:]:
        extra._element.getparent().remove(extra._element)


def _doc_has_token(doc, token: str) -> bool:
    for p in doc.paragraphs:
        if token in p.text:
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if token in cell.text:
                    return True
    return False


def patch_optional_terms_row(doc) -> bool:
    """Append the Optional Terms row to the Dispute-Resolution appendix
    table. Returns True if a row was added, False if it was already there."""
    if _doc_has_token(doc, _ROW_TOKEN):
        return False

    target_table = None
    for table in doc.tables:
        for row in table.rows:
            if row.cells[0].text.strip().lower().startswith("dispute resolution"):
                target_table = table
                break
        if target_table is not None:
            break
    if target_table is None:
        raise RuntimeError(
            "Could not find the 'Dispute Resolution (Jurisdiction)' appendix "
            "row to append the Optional Terms row after."
        )

    last_row = target_table.rows[-1]
    new_tr = deepcopy(last_row._tr)
    target_table._tbl.append(new_tr)

    new_row = target_table.rows[-1]
    _set_cell_text(new_row.cells[0], _ROW_LABEL)
    if len(new_row.cells) > 2:
        _set_cell_text(new_row.cells[1], "")  # no spec clause number
        _set_cell_text(new_row.cells[2], _ROW_TOKEN)
    else:
        _set_cell_text(new_row.cells[-1], _ROW_TOKEN)
    return True


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    added = patch_optional_terms_row(doc)
    if not added:
        print("{{C15}} already present — no change.")
        return

    doc.save(str(MASTER_DOCX))

    # Confirm the token landed.
    doc = Document(str(MASTER_DOCX))
    if not _doc_has_token(doc, _ROW_TOKEN):
        raise SystemExit("Patch ran but {{C15}} is still missing.")
    print(f"Appended '{_ROW_LABEL}' row with {_ROW_TOKEN} to the page-7 appendix table.")


if __name__ == "__main__":
    main()
