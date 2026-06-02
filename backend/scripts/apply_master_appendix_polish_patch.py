"""Apply appendix-polish patches to the master docx.

Changes applied (all idempotent):
1. Remove 'Maximum Liquidated Damages' row  (Table 4, row containing {{A21_DISPLAY}}).
2. Rate Of Liquidated Damages value cell     (Table 4): {{C11}}  -> AED {{C11}} per day
3. The Subcontract Price value cell          (Table 3): {{F08}}  -> AED {{F08}}
4. Insurance Policies value cell             (Table 4): {{A22}}  -> {{A22}} days
   + add missing <w:w val='105'/> run-property so the cell renders in the same
     character-width scaling as every other row in the table.
5. Insurance Policies label cell             (Table 4): rebuild all runs to match
   Rate Of Liquidated Damages styling (remove italic, sz 23→20, add w:w=105,
   fix pPr spacing and indent).

Run::

    python backend/scripts/apply_master_appendix_polish_patch.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"


def _set_value_cell(cell, new_text: str) -> None:
    """Replace value-cell content with new_text, preserving the first run's rPr."""
    para = cell.paragraphs[0]
    # Remove extra paragraphs
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if para.runs:
        para.runs[0].text = new_text
        for extra_run in para.runs[1:]:
            extra_run._r.getparent().remove(extra_run._r)
    else:
        para.add_run(new_text)


def _fix_insurance_label_cell(label_cell, reference_label_cell) -> bool:
    """Rebuild the Insurance Policies label cell to match the style of the
    Rate Of Liquidated Damages label cell (remove italic, correct font size,
    fix paragraph spacing/indent).

    Returns True if changes were made.
    """
    from copy import deepcopy

    # Already fixed if no italic runs present
    for para in label_cell.paragraphs:
        for run in para.runs:
            if run.italic:
                break
        else:
            continue
        break
    else:
        return False  # no italic found — already clean

    # ── Copy pPr from the reference row ──────────────────────────────────────
    ref_para = reference_label_cell.paragraphs[0]
    tgt_para = label_cell.paragraphs[0]
    ref_ppr = ref_para._element.find(qn("w:pPr"))
    tgt_ppr = tgt_para._element.find(qn("w:pPr"))

    if ref_ppr is not None:
        new_ppr = deepcopy(ref_ppr)
        if tgt_ppr is not None:
            tgt_para._element.replace(tgt_ppr, new_ppr)
        else:
            tgt_para._element.insert(0, new_ppr)

    # ── Collect full text then replace all runs with a single clean run ───────
    full_text = "".join(r.text or "" for r in tgt_para.runs)

    # Remove all existing runs
    for run in list(tgt_para.runs):
        run._r.getparent().remove(run._r)

    # Build a new run copying rPr from the first reference run (w:w=105, sz=20)
    ref_runs = ref_para.runs
    base_rpr = deepcopy(ref_runs[0]._r.find(qn("w:rPr"))) if ref_runs else None

    new_r = OxmlElement("w:r")
    if base_rpr is not None:
        # Strip any spacing tweak — plain font with no letter-spacing delta
        for sp in base_rpr.findall(qn("w:spacing")):
            base_rpr.remove(sp)
        new_r.append(base_rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = full_text
    new_r.append(t)
    tgt_para._element.append(new_r)

    # Remove extra paragraphs (the cell may have had wrapped paras)
    for extra in label_cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    return True


def _ensure_w_w(cell, val: str = "105") -> None:
    """Ensure the first run in cell has <w:w val='{val}'/> in its rPr."""
    para = cell.paragraphs[0]
    if not para.runs:
        return
    run = para.runs[0]
    r_el = run._r
    rpr = r_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r_el.insert(0, rpr)
    # Remove any existing w:w element then re-add
    for existing in rpr.findall(qn("w:w")):
        rpr.remove(existing)
    w_el = OxmlElement("w:w")
    w_el.set(qn("w:val"), val)
    # Insert after w:spacing (first child) if present, else append
    spacing = rpr.find(qn("w:spacing"))
    if spacing is not None:
        idx = list(rpr).index(spacing)
        rpr.insert(idx + 1, w_el)
    else:
        rpr.insert(0, w_el)


def patch(doc) -> dict[str, bool]:
    """Apply all patches. Returns dict of change flags."""
    changes: dict[str, bool] = {}

    # ── Table 3 ── Subcontract Price (F08) ───────────────────────────────────
    t3 = doc.tables[3]
    for row in t3.rows:
        val_text = row.cells[-1].text.strip()
        if val_text == "{{F08}}":
            _set_value_cell(row.cells[-1], "AED {{F08}}")
            changes["t3_f08_aed"] = True
            break
    else:
        changes["t3_f08_aed"] = False  # already done or not found

    # ── Table 4 ──────────────────────────────────────────────────────────────
    t4 = doc.tables[4]

    # 1. Remove Maximum Liquidated Damages row ({{A21_DISPLAY}})
    a21_row = None
    for row in t4.rows:
        if "A21_DISPLAY" in row.cells[-1].text or "{{A21" in row.cells[-1].text:
            a21_row = row
            break
    if a21_row is not None:
        a21_row._tr.getparent().remove(a21_row._tr)
        changes["remove_a21_row"] = True
    else:
        changes["remove_a21_row"] = False

    # Re-iterate after possible deletion
    rate_ld_row = None
    insurance_row = None
    for row in t4.rows:
        val_text = row.cells[-1].text.strip()

        # 2. Rate Of Liquidated Damages: {{C11}} → AED {{C11}} per day
        if val_text in ("{{C11}}", "AED {{C11}} per day"):
            if val_text == "{{C11}}":
                _set_value_cell(row.cells[-1], "AED {{C11}} per day")
                changes["c11_per_day"] = True
            rate_ld_row = row

        # 3. Insurance Policies: {{A22}} → {{A22}} days + fix rPr w:w
        elif "{{A22}}" in val_text:
            if val_text == "{{A22}}":
                _set_value_cell(row.cells[-1], "{{A22}} days")
                _ensure_w_w(row.cells[-1])
                changes["a22_days_fix"] = True
            insurance_row = row

    # 4. Fix Insurance Policies label cell styling (italic / wrong font size)
    if rate_ld_row is not None and insurance_row is not None:
        fixed = _fix_insurance_label_cell(insurance_row.cells[0], rate_ld_row.cells[0])
        changes["a22_label_fix"] = fixed
    else:
        changes["a22_label_fix"] = False

    return changes


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    results = patch(doc)

    if results.get("t3_f08_aed"):
        print("Table 3: 'The Subcontract Price' value cell updated → AED {{F08}}")
    if results.get("remove_a21_row"):
        print("Table 4: 'Maximum Liquidated Damages' row removed.")
    if results.get("c11_per_day"):
        print("Table 4: 'Rate Of Liquidated Damages' value cell updated → AED {{C11}} per day")
    if results.get("a22_days_fix"):
        print("Table 4: 'Insurance Policies' value cell updated → {{A22}} days (+ rPr fix)")
    if results.get("a22_label_fix"):
        print("Table 4: 'Insurance Policies' label cell rebuilt — italic removed, sz 23→20, styling matched to Rate Of Liquidated Damages.")

    applied = any(results.values())
    if not applied:
        print("All patches already applied — no change.")
        return

    doc.save(str(MASTER_DOCX))

    # Verify
    doc2 = Document(str(MASTER_DOCX))
    for t in doc2.tables:
        for row in t.rows:
            if "A21_DISPLAY" in row.cells[-1].text:
                raise SystemExit("Verification failed: {{A21_DISPLAY}} still present.")
    print("Verification passed.")


if __name__ == "__main__":
    main()
