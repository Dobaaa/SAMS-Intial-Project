"""Rewrite the Advance Payment, Progress Payment, and Retention release
clauses with BGCC's fixed PDC / Engineer's Certificate wording.

Phase 2 Package G (req 12, 13, 14). Client-confirmed exact wording, replacing
the free-text {{C04}}/{{C05}}/{{C06}}/{{C07}} tokens with fixed sentences —
these terms are now standard policy, not admin-entered per agreement:

- C04 (Advance Payment release condition, para ~271): drops the old
  "{{C04}}" day-count entirely; release is now conditioned on Bank
  Guarantee submission + engineering-works approval only.
- C05 (Interim/Progress Payment, para ~311): drops the old "back-to-back"
  trigger (BGCC only paying once paid by the Employer) entirely per
  client instruction — replaced by a fixed 60-day PDC / 15-day
  Engineer's-Certificate mechanism.
- C06/C07 (1st/2nd retention instalments, paras ~321/328): the old
  itemized bullet lists (Employer payment receipt, defect rectification,
  close-out docs, final account certification / DLP expiry, Final
  Completion Certificate, Employer retention receipt) are DELETED —
  replaced by one self-contained sentence each, per client instruction to
  use the wording as given.

C04/C05/C06/C07 fields become orphaned (no longer referenced by any token
in the master) — left as-is in seed_fields.py/MasterField, not hidden or
removed, matching the lesson from the A15 incident: don't proactively
remove/hide data-model fields beyond what was actually asked.

Idempotent: if none of {{C04}}/{{C05}}/{{C06}}/{{C07}} are present, exits
without changes.

Run::

    python backend/scripts/apply_master_payment_retention_wording_patch.py
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

BACKEND_DIR = Path(__file__).resolve().parents[1]
MASTER_DOCX = BACKEND_DIR / "masters" / "sca_master_v1.docx"
BACKUP = MASTER_DOCX.with_suffix(".docx.pre-payment-retention-wording.bak")

TOKENS = ("{{C04}}", "{{C05}}", "{{C06}}", "{{C07}}")

C04_TEXT = (
    "The advance payment shall be released upon submission of the required "
    "Bank Guarantee and approval of the engineering works."
)
C05_TEXT = (
    "Progress payments shall be released by 60-day Post-Dated Cheque (PDC). "
    "Payment shall be made within 15 days from the date of approval of the "
    "Engineer's Certificate, with the PDC dated 60 days from the invoice date."
)
C06_TEXT = (
    "The first 50% of the retention shall be released by 60-day Post-Dated "
    "Cheque (PDC) from the invoice date upon issuance of the Taking-Over "
    "Certificate (TOC) and within 15 days from the date of approval of the "
    "Engineer's Certificate."
)
C07_TEXT = (
    "The remaining 50% of the retention shall be released by 60-day "
    "Post-Dated Cheque (PDC) from the invoice date upon completion of the "
    "Defects Liability Period (365 days) and within 15 days from the date "
    "of approval of the Engineer's Certificate."
)


def _doc_text(doc) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def _set_paragraph_text(para, text: str) -> None:
    """Replace a paragraph's visible text via run 0 (clearing any other
    runs), preserving paragraph-level formatting. Matches the pattern
    already used for the Commencement Date clause-4.3 rewrite."""
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def _find_paragraph(doc, token: str):
    for p in doc.paragraphs:
        if token in p.text:
            return p
    raise RuntimeError(f"Could not find a paragraph containing {token}")


def _delete_paragraph(para) -> None:
    para._element.getparent().remove(para._element)


def main() -> None:
    if not MASTER_DOCX.exists():
        raise SystemExit(f"Master docx not found at {MASTER_DOCX}")

    doc = Document(str(MASTER_DOCX))
    full_text = _doc_text(doc)
    present = [t for t in TOKENS if t in full_text]

    if not present:
        print("Payment/retention wording already patched — no change.")
        return
    if len(present) != len(TOKENS):
        raise SystemExit(
            f"Partial patch state detected — found {present}, expected all of "
            f"{list(TOKENS)}. Investigate before re-running."
        )

    # C04 — Advance Payment release condition.
    _set_paragraph_text(_find_paragraph(doc, "{{C04}}"), C04_TEXT)

    # C05 — Progress Payments.
    _set_paragraph_text(_find_paragraph(doc, "{{C05}}"), C05_TEXT)

    # C06 — 1st retention instalment: rewrite the intro sentence, then
    # delete the four bullet sub-paragraphs that followed it (322-326 in
    # the original layout) — the new sentence is self-contained.
    c06_para = _find_paragraph(doc, "{{C06}}")
    c06_bullets = []
    node = c06_para._p.getnext()
    while node is not None and node.tag == c06_para._p.tag:
        from docx.text.paragraph import Paragraph

        candidate = Paragraph(node, c06_para._parent)
        if "{{C07}}" in candidate.text or not candidate.text.strip():
            break
        c06_bullets.append(candidate)
        node = node.getnext()
    _set_paragraph_text(c06_para, C06_TEXT)
    for b in c06_bullets:
        _delete_paragraph(b)

    # C07 — 2nd retention instalment: same treatment.
    c07_para = _find_paragraph(doc, "{{C07}}")
    c07_bullets = []
    node = c07_para._p.getnext()
    while node is not None and node.tag == c07_para._p.tag:
        from docx.text.paragraph import Paragraph

        candidate = Paragraph(node, c07_para._parent)
        if not candidate.text.strip():
            break
        c07_bullets.append(candidate)
        node = node.getnext()
    _set_paragraph_text(c07_para, C07_TEXT)
    for b in c07_bullets:
        _delete_paragraph(b)

    if not BACKUP.exists():
        shutil.copy2(MASTER_DOCX, BACKUP)

    doc.save(str(MASTER_DOCX))

    doc2 = Document(str(MASTER_DOCX))
    text2 = _doc_text(doc2)
    for t in TOKENS:
        if t in text2:
            raise SystemExit(f"Patch ran but {t} is still present.")
    for expected in (C04_TEXT, C05_TEXT, C06_TEXT, C07_TEXT):
        if expected not in text2:
            raise SystemExit(f"Patch ran but expected text not found: {expected[:60]}...")
    print("Advance Payment / Progress Payment / Retention wording updated.")


if __name__ == "__main__":
    main()
